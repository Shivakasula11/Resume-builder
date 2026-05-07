"""
AI Resume Analyzer v3.0 — Real JD-Based Analysis
==================================================
- Store real JDs in SQLite per role
- Extract skills from each JD using GPT
- Find 70%+ common skills across all JDs for a role
- Compare resume against REAL common skills
- Score = how many real JD-common skills your resume has

Deps: pip install openai python-dotenv reportlab
"""

import os
import re
import json
import hashlib
import sqlite3
from typing import List, Dict, Optional
from pathlib import Path
from datetime import datetime
from collections import Counter

from dotenv import load_dotenv
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
CACHE_DIR = Path("uploads/analysis_cache")
CACHE_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = Path("uploads/jd_database.db")
DB_PATH.parent.mkdir(parents=True, exist_ok=True)


def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS job_descriptions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            role TEXT NOT NULL,
            level TEXT NOT NULL DEFAULT 'fresher',
            jd_text TEXT NOT NULL,
            company TEXT DEFAULT '',
            source TEXT DEFAULT '',
            extracted_skills TEXT DEFAULT '[]',
            skill_count INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_role_level ON job_descriptions(role, level)")
    conn.commit()
    conn.close()

init_db()


def call_openai(system_prompt, user_prompt, max_tokens=4000):
    from openai import OpenAI
    client = OpenAI(api_key=OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        max_tokens=max_tokens, temperature=0.2,
    )
    return response.choices[0].message.content.strip()


def parse_json_response(text):
    text = text.strip()
    text = re.sub(r'^```json\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    return json.loads(text.strip())


# ═══════════════════════════════════════════
#  PROMPTS
# ═══════════════════════════════════════════
JD_EXTRACT_PROMPT = """Extract ONLY the technical skills, technologies, frameworks, libraries, and domain expertise that this job description EXPLICITLY requires or mentions.

RULES:
1. Only extract skills WRITTEN in the JD text. Do NOT add skills the JD does not mention.
2. DO NOT extract: soft skills, IDEs, Git, JIRA, Agile/Scrum, OS names, generic tools
3. Normalize: "Python 3" -> "Python", "Amazon Web Services" -> "AWS"
4. If JD says "ML frameworks like TensorFlow and PyTorch", extract: TensorFlow, PyTorch, Machine Learning

Return ONLY valid JSON array:
[{"skill": "Python", "category": "programming"}, ...]

Categories: programming, framework, core_concept, database, cloud

JD TEXT:
{jd_text}"""

RESUME_EXTRACT_PROMPT = """Extract ONLY technical skills, technologies, libraries, frameworks, and domain expertise from this resume.

RULES:
1. Look everywhere: summary, skills, experience bullets, projects, education, certifications
2. If tech mentioned in project ("Built with Django and SQLite"), extract BOTH
3. Extract domain knowledge from context ("emotion detection from audio" = Speech Processing)
4. Extract specific techniques: MFCC, FAISS, RAG, SBERT, etc.

DO NOT EXTRACT: Soft skills, IDEs/editors, Git/GitHub, JIRA, Postman, OS names, Agile/Scrum, Debugging, Testing, Documentation, OOP, Data Structures

Return ONLY valid JSON array:
[{"skill": "Python", "category": "programming"}, ...]

Categories: programming, framework, core_concept, database, cloud

RESUME TEXT:
{resume_text}"""


# ═══════════════════════════════════════════
#  ALIASES & NORMALIZATION
# ═══════════════════════════════════════════
SKILL_ALIASES = {
    "python": ["python3", "python 3", "python programming"],
    "javascript": ["js", "es6", "ecmascript"],
    "typescript": ["ts"],
    "java": ["core java", "java programming", "java se"],
    "c++": ["cpp", "c plus plus"],
    "c#": ["csharp", "c sharp"],
    "html": ["html5", "html 5"],
    "css": ["css3", "css 3"],
    "machine learning": ["ml", "machine-learning", "ml algorithms"],
    "deep learning": ["dl", "deep-learning", "deep neural networks"],
    "artificial intelligence": ["ai"],
    "natural language processing": ["nlp", "text processing", "text mining"],
    "computer vision": ["cv", "image processing", "image recognition"],
    "amazon web services": ["aws", "amazon cloud"],
    "google cloud platform": ["gcp", "google cloud"],
    "microsoft azure": ["azure"],
    "tensorflow": ["tf", "tensor flow"],
    "pytorch": ["torch"],
    "scikit-learn": ["sklearn", "scikit learn"],
    "sql": ["structured query language", "sql queries"],
    "mysql": ["my sql", "my-sql"],
    "sqlite": ["sqlite3"],
    "nosql": ["no-sql"],
    "kubernetes": ["k8s"],
    "docker": ["containerization", "containers"],
    "large language models": ["llm", "llms", "large language model"],
    "generative ai": ["gen ai", "genai"],
    "rag": ["retrieval augmented generation", "rag pipeline"],
    "langchain": ["lang chain"],
    "hugging face": ["huggingface", "hf transformers"],
    "vector database": ["vector databases", "vector db", "vector store"],
    "faiss": ["facebook ai similarity search"],
    "react": ["reactjs", "react.js"],
    "node.js": ["nodejs", "node"],
    "express": ["expressjs", "express.js"],
    "mongodb": ["mongo"],
    "postgresql": ["postgres"],
    "fastapi": ["fast api", "fast-api"],
    "django": ["django rest framework", "drf", "django rest"],
    "flask": ["flask api"],
    "opencv": ["open cv", "cv2"],
    "keras": ["keras api"],
    "pandas": ["pandas library"],
    "numpy": ["numpy library"],
    "speech processing": ["speech recognition", "speech emotion", "audio processing"],
    "face recognition": ["facial recognition", "face detection"],
    "sbert": ["sentence bert", "sentence-bert", "sentence transformers"],
    "mfcc": ["mel frequency cepstral coefficients"],
    "lstm": ["long short-term memory", "long short term memory"],
    "cnn": ["convolutional neural network", "convolutional neural networks"],
    "rnn": ["recurrent neural network", "recurrent neural networks"],
    "data science": ["data analytics", "data analysis"],
    "nlp": ["natural language processing", "text processing"],
    "power bi": ["powerbi"],
    "tableau": ["tableau desktop"],
    "apache spark": ["spark", "pyspark"],
    "apache kafka": ["kafka"],
    "apache airflow": ["airflow"],
    "mlflow": ["ml flow"],
    "kubeflow": ["kube flow"],
    "terraform": ["terraform iac"],
}

def normalize_skill(name):
    n = name.strip().title()
    for old, new in [("Aws","AWS"),("Gcp","GCP"),("Sql","SQL"),("Nosql","NoSQL"),
                     ("Ci/Cd","CI/CD"),("Ml","ML"),("Ai","AI"),("Nlp","NLP"),
                     ("Api","API"),("Sdk","SDK"),("Llm","LLM"),("Gpu","GPU"),
                     ("Dl","DL"),("Cv","CV"),("Devops","DevOps"),("Node.Js","Node.js"),
                     ("React.Js","React.js"),("Html","HTML"),("Css","CSS"),("Dsa","DSA"),
                     ("Mlops","MLOps"),("Mlflow","MLflow")]:
        n = n.replace(old, new)
    return n

def _canonicalize(skill):
    s = skill.lower().strip()
    for canonical, aliases in SKILL_ALIASES.items():
        if s == canonical or s in aliases:
            return canonical
    return s


BLACKLISTED_SKILLS = {
    "vs code","visual studio code","pycharm","intellij","sublime text","atom",
    "jupyter notebook","jupyter","jupyter lab","google colab","colab",
    "visual studio","eclipse","netbeans","vim","emacs",
    "jira","slack","trello","confluence","notion","postman","swagger",
    "microsoft office","ms office","word","excel","powerpoint",
    "communication","communication skills","teamwork","team collaboration",
    "leadership","problem solving","problem-solving","adaptability",
    "time management","critical thinking","creativity","attention to detail",
    "self motivated","work ethic","interpersonal skills","presentation skills",
    "analytical thinking","decision making","collaboration","team player",
    "effective communication","verbal communication","written communication",
    "quick learner","fast learner","proactive","detail oriented",
    "dedicated","passionate","enthusiastic","hardworking",
    "innovative","strategic thinking","mentoring","coaching",
    "git","github","gitlab","bitbucket","version control",
    "linux","windows","macos","ubuntu","centos","unix",
    "agile","scrum","waterfall","kanban",
    "debugging","testing","unit testing","documentation","code review",
    "clean code","oop","object oriented programming",
    "data structures","algorithms","data structures and algorithms","dsa",
}

def filter_blacklisted(skills):
    return [s for s in skills
            if s.get("skill","").strip().lower() not in BLACKLISTED_SKILLS
            and s.get("category","").lower() not in ("soft_skill","tool","methodology")]


# ═══════════════════════════════════════════
#  JD MANAGEMENT (SQLite)
# ═══════════════════════════════════════════
def add_job_description(role, level, jd_text, company="", source=""):
    role_normalized = role.strip().lower()
    prompt = JD_EXTRACT_PROMPT.replace("{jd_text}", jd_text[:6000])
    response = call_openai(
        "Extract only technical skills from job descriptions. No soft skills, no generic tools. Return only valid JSON arrays.",
        prompt
    )
    raw_skills = parse_json_response(response)

    seen = set()
    skills = []
    for s in raw_skills:
        if isinstance(s, dict) and "skill" in s:
            normalized = normalize_skill(s["skill"])
            canonical = _canonicalize(normalized)
            if canonical not in seen:
                seen.add(canonical)
                s["skill"] = normalized
                s["canonical"] = canonical
                skills.append(s)
    skills = filter_blacklisted(skills)

    conn = get_db()
    cursor = conn.execute(
        "INSERT INTO job_descriptions (role, level, jd_text, company, source, extracted_skills, skill_count) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (role_normalized, level.lower(), jd_text.strip(), company.strip(), source.strip(),
         json.dumps(skills), len(skills))
    )
    conn.commit()
    jd_id = cursor.lastrowid
    conn.close()

    return {
        "id": str(jd_id), "role": role_normalized, "level": level,
        "company": company, "skills_extracted": len(skills),
        "skills": [s["skill"] for s in skills],
    }


def get_jds_for_role(role, level="fresher"):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM job_descriptions WHERE role=? AND level=? ORDER BY created_at DESC",
        (role.strip().lower(), level.lower())
    ).fetchall()
    conn.close()

    results = []
    for row in rows:
        results.append({
            "_id": str(row["id"]),
            "role": row["role"],
            "level": row["level"],
            "jd_text": row["jd_text"],
            "company": row["company"],
            "source": row["source"],
            "extracted_skills": json.loads(row["extracted_skills"]),
            "skill_count": row["skill_count"],
            "created_at": row["created_at"],
        })
    return results


def delete_job_description(jd_id):
    conn = get_db()
    result = conn.execute("DELETE FROM job_descriptions WHERE id=?", (int(jd_id),))
    conn.commit()
    deleted = result.rowcount > 0
    conn.close()
    return deleted


def get_all_roles():
    conn = get_db()
    rows = conn.execute("""
        SELECT role, level, COUNT(*) as count, MAX(created_at) as latest
        FROM job_descriptions
        GROUP BY role, level
        ORDER BY role, level
    """).fetchall()
    conn.close()

    return [{
        "role": row["role"], "level": row["level"],
        "jd_count": row["count"], "latest_added": row["latest"] or "",
    } for row in rows]


def get_jd_count():
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM job_descriptions").fetchone()[0]
    conn.close()
    return count


# ═══════════════════════════════════════════
#  ANALYZE COMMON SKILLS FROM REAL JDs
# ═══════════════════════════════════════════
def analyze_common_skills(role, level="fresher", min_frequency_pct=70.0):
    jds = get_jds_for_role(role, level)
    total_jds = len(jds)

    if total_jds == 0:
        return {"error": f"No JDs found for '{role}' ({level}). Add JDs first via the JD Manager.",
                "total_jds": 0, "skills": []}

    skill_counter = Counter()
    skill_details = {}

    for jd in jds:
        jd_seen = set()
        for s in jd.get("extracted_skills", []):
            canonical = s.get("canonical", _canonicalize(s["skill"]))
            if canonical not in jd_seen:
                jd_seen.add(canonical)
                skill_counter[canonical] += 1
                if canonical not in skill_details:
                    skill_details[canonical] = {
                        "skill": normalize_skill(s["skill"]),
                        "category": s.get("category", "core_concept"),
                    }

    min_count = (min_frequency_pct / 100.0) * total_jds
    common_skills = []
    for canonical, count in skill_counter.most_common():
        freq = round((count / total_jds) * 100, 1)
        if count >= min_count:
            info = skill_details[canonical]
            common_skills.append({
                "skill": info["skill"], "canonical": canonical,
                "category": info["category"], "frequency": freq,
                "jd_count": count,
                "importance": "required" if freq >= 85 else "important",
            })

    common_skills.sort(key=lambda x: x["frequency"], reverse=True)
    return {
        "role": role, "level": level, "total_jds": total_jds,
        "min_frequency_pct": min_frequency_pct,
        "total_common_skills": len(common_skills),
        "skills": common_skills, "all_skills_found": len(skill_counter),
    }


# ═══════════════════════════════════════════
#  RESUME EXTRACTION
# ═══════════════════════════════════════════
def extract_text_from_docx(docx_path):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text: parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text: parts.append(text)
    body = doc.element.body
    for txbx in body.iter(qn('w:txbxContent')):
        text = "".join((t.text or "") for t in txbx.iter(qn('w:t'))).strip()
        if text: parts.append(text)
    return "\n".join(parts)


def extract_resume_skills(docx_path):
    text = extract_text_from_docx(docx_path)
    if not text or len(text.strip()) < 20: return []
    if len(text) > 8000: text = text[:8000]

    prompt = RESUME_EXTRACT_PROMPT.replace("{resume_text}", text)
    response = call_openai(
        "Extract technical skills from resumes. Only real technical skills. Return only valid JSON arrays.",
        prompt
    )
    skills = parse_json_response(response)

    seen = set()
    unique = []
    for s in skills:
        if isinstance(s, dict) and "skill" in s:
            normalized = normalize_skill(s["skill"])
            key = normalized.lower()
            if key not in seen:
                seen.add(key)
                s["skill"] = normalized
                s["canonical"] = _canonicalize(normalized)
                unique.append(s)
    return filter_blacklisted(unique)


# ═══════════════════════════════════════════
#  SKILL MATCHING
# ═══════════════════════════════════════════
def skills_match(jd_skill, resume_skills_lower):
    jd_lower = jd_skill.lower().strip()
    jd_canonical = _canonicalize(jd_skill)

    if jd_lower in resume_skills_lower:
        return True, jd_skill

    for rs in resume_skills_lower:
        if _canonicalize(rs) == jd_canonical:
            return True, rs.title()

    for canonical, aliases in SKILL_ALIASES.items():
        all_forms = [canonical] + aliases
        if jd_lower in all_forms:
            for form in all_forms:
                if form in resume_skills_lower:
                    return True, form.title()

    for rs in resume_skills_lower:
        if len(jd_lower) >= 3 and len(rs) >= 3:
            if jd_lower in rs or rs in jd_lower:
                return True, rs.title()

    jd_words = set(jd_lower.split())
    if len(jd_words) >= 2:
        for rs in resume_skills_lower:
            rs_words = set(rs.split())
            if len(rs_words) >= 2:
                overlap = jd_words & rs_words
                if len(overlap) >= len(jd_words) * 0.6:
                    return True, rs.title()

    return False, None


# ═══════════════════════════════════════════
#  ATS COMPARISON — Resume vs Each JD
# ═══════════════════════════════════════════
JD_RESUME_COMPARE_PROMPT = """Compare this resume against this job description and give a match score.

Analyze how well this candidate fits this specific job by checking:
1. Does the resume have the technical skills the JD asks for?
2. Does the candidate have relevant work experience or projects in the same area?
3. Does the candidate's domain knowledge align with what the job needs?
4. Does the candidate meet the education, certification, or years-of-experience requirements?

Be honest and proportional:
- Only count skills that genuinely match (Python is not ABAP, MySQL is not SAP HANA)
- If the JD asks for 10 skills and resume has 3, the score should reflect ~30% skill match
- Education alone should not inflate the score beyond 5%
- Soft skills should not count toward the score
- Don't give courtesy points for unrelated experience

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}

Return ONLY valid JSON:
{{
  "score": 75,
  "matched_skills": ["Python", "TensorFlow", "NLP"],
  "missing_skills": ["Docker", "Kubernetes"],
  "strengths": "Strong ML project portfolio with hands-on TensorFlow experience",
  "gaps": "No deployment/MLOps experience",
  "verdict": "Good fit for junior ML role, needs cloud skills for senior positions"
}}

Rules:
- score: 0-100 integer. Proportional to real matches.
- matched_skills: only skills from JD genuinely found in resume. Empty array [] if none.
- missing_skills: important skills from JD not in resume (max 10)
- strengths: 1 sentence on what actually matches
- gaps: 1 sentence on what's missing
- verdict: 1 sentence overall fit assessment"""


def _evaluate_single_jd(jd, resume_text):
    """Evaluate resume against a single JD using GPT."""
    try:
        prompt = JD_RESUME_COMPARE_PROMPT.replace("{jd_text}", jd["jd_text"][:4000]).replace("{resume_text}", resume_text[:4000])
        response = call_openai(
            "Compare the resume against the job description honestly. Score based on real skill matches. Return only valid JSON.",
            prompt
        )
        result = parse_json_response(response)
        return {
            "jd_id": jd["_id"],
            "company": jd.get("company", ""),
            "score": max(0, min(100, int(result.get("score", 0)))),
            "matched_skills": result.get("matched_skills", []),
            "missing_skills": result.get("missing_skills", []),
            "strengths": result.get("strengths", ""),
            "gaps": result.get("gaps", ""),
            "verdict": result.get("verdict", ""),
            "jd_preview": jd["jd_text"][:150] + "...",
        }
    except Exception as e:
        print(f"[ATS] Error evaluating JD {jd.get('_id','?')}: {e}")
        return {
            "jd_id": jd.get("_id", ""), "company": jd.get("company", ""),
            "score": 0, "matched_skills": [], "missing_skills": [],
            "strengths": "", "gaps": "", "verdict": f"Error: {str(e)[:80]}",
            "jd_preview": jd.get("jd_text", "")[:150] + "...",
        }


def compare_resume_to_role(session_id, docx_path, role, level="fresher", force_refresh=False):
    from concurrent.futures import ThreadPoolExecutor, as_completed

    # Step 1: Get all JDs
    jds = get_jds_for_role(role, level)
    if not jds:
        return {"error": f"No JDs found for '{role}' ({level}). Add JDs first via the JD Manager."}

    # Step 2: Extract resume text
    resume_text = extract_text_from_docx(docx_path)
    if not resume_text or len(resume_text.strip()) < 20:
        return {"error": "Could not extract text from resume."}
    if len(resume_text) > 5000:
        resume_text = resume_text[:5000]

    # Step 3: Evaluate against each JD in parallel
    jd_results = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(_evaluate_single_jd, jd, resume_text): jd for jd in jds}
        for future in as_completed(futures):
            jd_results.append(future.result())

    # Sort by score descending
    jd_results.sort(key=lambda x: x["score"], reverse=True)

    # Step 4: Calculate stats
    scores = [r["score"] for r in jd_results]
    mean_score = round(sum(scores) / len(scores), 1) if scores else 0
    max_score = max(scores) if scores else 0
    min_score = min(scores) if scores else 0

    # Step 5: Aggregate common matched & missing skills
    all_matched = Counter()
    all_missing = Counter()
    for r in jd_results:
        for s in r.get("matched_skills", []):
            all_matched[s] += 1
        for s in r.get("missing_skills", []):
            all_missing[s] += 1

    total_jds = len(jds)
    common_matched = [{"skill": s, "count": c, "pct": round(c/total_jds*100)}
                      for s, c in all_matched.most_common(20)]
    common_missing = [{"skill": s, "count": c, "pct": round(c/total_jds*100)}
                      for s, c in all_missing.most_common(15)]

    # Score distribution
    dist = {"90_100": 0, "70_89": 0, "50_69": 0, "below_50": 0}
    for s in scores:
        if s >= 90: dist["90_100"] += 1
        elif s >= 70: dist["70_89"] += 1
        elif s >= 50: dist["50_69"] += 1
        else: dist["below_50"] += 1

    comparison = {
        "session_id": session_id, "role": role, "level": level,
        "overall_score": mean_score, "max_score": max_score, "min_score": min_score,
        "total_jds": total_jds,
        "jd_results": jd_results,
        "score_distribution": dist,
        "common_matched": common_matched,
        "common_missing": common_missing,
        "compared_at": datetime.now().isoformat(),
    }

    ck = hashlib.md5(f"{role.strip().lower()}_{level}".encode()).hexdigest()
    rf = CACHE_DIR / f"comparison_{session_id}_{ck}.json"
    rf.write_text(json.dumps(comparison, indent=2, default=str), encoding="utf-8")
    return comparison


# ═══════════════════════════════════════════
#  PDF REPORT
# ═══════════════════════════════════════════
def generate_comparison_pdf(comparison, output_path):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_CENTER

        white = HexColor("#ffffff"); gray = HexColor("#d1d5db")
        primary = HexColor("#4f46e5"); success = HexColor("#059669")
        danger = HexColor("#dc2626"); warning = HexColor("#d97706")
        light_bg = HexColor("#f9fafb"); dark = HexColor("#111827")

        doc = SimpleDocTemplate(output_path, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        elements = []

        ts = ParagraphStyle('T', parent=styles['Title'], fontSize=22, textColor=dark, spaceAfter=4)
        h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, textColor=primary, spaceBefore=12, spaceAfter=6)
        bs = ParagraphStyle('B', parent=styles['Normal'], fontSize=10, leading=14, textColor=HexColor("#374151"))
        ss = ParagraphStyle('S', parent=styles['Normal'], fontSize=8, textColor=HexColor("#9ca3af"), alignment=TA_CENTER)

        score = comparison["overall_score"]
        total_jds = comparison.get("total_jds", 0)
        level = comparison.get("level", "fresher").title()
        sc = success if score >= 70 else warning if score >= 50 else danger

        elements.append(Paragraph(f"Resume Match Report - {comparison['role']}", ts))
        elements.append(Paragraph(f"Level: {level} | Compared against {total_jds} real JDs", bs))
        elements.append(Spacer(1, 6))

        # Score summary
        st = Table([[
            Paragraph(f'<font size="28" color="{sc.hexval()}">{score}%</font>', ParagraphStyle('sc', alignment=TA_CENTER)),
            Paragraph(f'<b>Average Match Score</b><br/>Best: {comparison.get("max_score",0)}% | '
                      f'Worst: {comparison.get("min_score",0)}% | JDs: {total_jds}', bs)
        ]], colWidths=[40*mm, 120*mm])
        st.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(0,0),HexColor("#f1f5f9")),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('LEFTPADDING',(0,0),(-1,-1),10),('RIGHTPADDING',(0,0),(-1,-1),10),
            ('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8),
        ]))
        elements.append(st); elements.append(Spacer(1, 12))

        # Per-JD scores
        jd_results = comparison.get("jd_results", [])
        if jd_results:
            elements.append(Paragraph("Per-JD Scores", h2))
            jd_data = [["#", "Company", "Score", "Verdict"]]
            for i, r in enumerate(jd_results, 1):
                sc_color = "green" if r["score"] >= 70 else "orange" if r["score"] >= 50 else "red"
                jd_data.append([str(i), r.get("company","") or f"JD #{i}",
                               f'{r["score"]}%', r.get("verdict","")[:80]])
            jt = Table(jd_data, colWidths=[10*mm, 35*mm, 20*mm, 95*mm])
            jt.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),HexColor("#4b5563")),('TEXTCOLOR',(0,0),(-1,0),white),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),8),
                ('ALIGN',(0,0),(2,-1),'CENTER'),('GRID',(0,0),(-1,-1),0.5,gray),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[white,light_bg]),
                ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
            ]))
            elements.append(jt); elements.append(Spacer(1, 16))

        # Common matched
        cm = comparison.get("common_matched", [])
        if cm:
            elements.append(Paragraph("Most Matched Skills (across JDs)", h2))
            elements.append(Paragraph(", ".join(f"{s['skill']} ({s['pct']}%)" for s in cm[:15]), bs))
            elements.append(Spacer(1, 12))

        # Common missing
        cmi = comparison.get("common_missing", [])
        if cmi:
            elements.append(Paragraph("Most Missing Skills (learn these first)", h2))
            elements.append(Paragraph(", ".join(f"{s['skill']} ({s['pct']}%)" for s in cmi[:10]), bs))
            elements.append(Spacer(1, 12))

        # Recommendation
        elements.append(Paragraph("Recommendation", h2))
        if score >= 80: elements.append(Paragraph("Strong match. Your resume fits most JDs well.", bs))
        elif score >= 60: elements.append(Paragraph("Moderate match. Focus on missing skills to improve.", bs))
        elif score >= 40: elements.append(Paragraph("Below average. Significant skill gaps need attention.", bs))
        else: elements.append(Paragraph("Low match. Major upskilling needed for this role.", bs))

        elements.append(Spacer(1, 30))
        elements.append(Paragraph(
            f"Generated by Resume Editor AI | {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
            f"Level: {level} | {total_jds} JDs analyzed", ss))
        doc.build(elements)
        return True
    except ImportError:
        print("[PDF] reportlab not installed"); return False
    except Exception as e:
        print(f"[PDF] Error: {e}"); import traceback; traceback.print_exc(); return False


def get_comparison_history(session_id):
    results = []
    for f in CACHE_DIR.glob(f"comparison_{session_id}_*.json"):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            results.append({
                "role": data.get("role",""), "level": data.get("level","fresher"),
                "overall_score": data.get("overall_score",0),
                "total_jds": data.get("total_jds",0),
                "compared_at": data.get("compared_at",""), "file": f.name,
            })
        except: pass
    return sorted(results, key=lambda x: x.get("compared_at",""), reverse=True)