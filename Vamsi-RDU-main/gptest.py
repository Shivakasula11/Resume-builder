"""
Test script — GPT-4o Vision PDF→DOCX conversion
=================================================
Install:
    pip install pymupdf openai python-docx pillow

Run:
    python test_gpt4o_pdf.py yourresume.pdf

Output:
    yourresume_gpt4o.docx — open in Word to check quality

Requires OPENAI_API_KEY in environment or .env file
"""

import sys
import os
import base64
import json
import time
from pathlib import Path

# Load .env if present
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
if not OPENAI_API_KEY:
    print("❌ OPENAI_API_KEY not set. Add it to .env or set as env variable.")
    sys.exit(1)


# ── Step 1: PDF → Images using PyMuPDF ──────────────────────────────

def pdf_to_images(pdf_path: Path) -> list[bytes]:
    """Convert each PDF page to a PNG image (base64)."""
    import fitz
    doc = fitz.open(str(pdf_path))
    images = []
    for page in doc:
        # 2x zoom for better OCR quality
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        images.append(pix.tobytes("png"))
        print(f"  ✓ Page {page.number + 1}/{len(doc)} rendered")
    doc.close()
    return images


# ── Step 2: GPT-4o Vision → Structured JSON ─────────────────────────

SYSTEM_PROMPT = """You are a precise document parser. 
You will receive an image of a resume page and must extract its EXACT content and structure.

Return a JSON object with this structure:
{
  "sections": [
    {
      "type": "header",           // header | section_title | entry | bullet | text | table_row
      "content": "...",           // plain text content
      "bold_parts": ["word1"],    // list of words/phrases that are bold
      "italic_parts": ["word1"],  // list of words/phrases that are italic
      "font_size": "large",       // large | medium | small | tiny
      "alignment": "left",        // left | center | right
      "indent_level": 0,          // 0=no indent, 1=one level, 2=two levels
      "is_bullet": false,         // true if this is a bullet point
      "bullet_char": "-",         // the bullet character if is_bullet=true
      "columns": null             // if type=table_row, array of column texts
    }
  ]
}

Rules:
- Preserve EXACT text including punctuation, symbols, percentages, dates
- Detect bold text (usually job titles, company names, skills, percentages)
- Detect italic text (usually subtitles, technologies)
- Detect bullet points (lines starting with –, -, •, ▪ etc)
- Detect two-column rows (name on left, date/location on right)
- Return ONLY valid JSON, no markdown, no explanation
"""

def extract_structure_gpt4o(image_bytes: bytes, page_num: int) -> dict:
    """Send page image to GPT-4o and get structured JSON back."""
    import openai
    client = openai.OpenAI(api_key=OPENAI_API_KEY)

    b64 = base64.b64encode(image_bytes).decode()

    print(f"  📤 Sending page {page_num} to GPT-4o…")
    t0 = time.time()

    response = client.chat.completions.create(
        model="gpt-4o",
        max_tokens=4096,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{b64}",
                            "detail": "high"
                        }
                    },
                    {
                        "type": "text",
                        "text": "Extract the complete structure of this resume page as JSON."
                    }
                ]
            }
        ]
    )

    raw = response.choices[0].message.content.strip()
    print(f"  ✓ GPT-4o responded in {time.time()-t0:.1f}s")

    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  ⚠ JSON parse error: {e}")
        print(f"  Raw response: {raw[:300]}")
        return {"sections": []}


# ── Step 3: Structured JSON → DOCX ──────────────────────────────────

def build_docx(all_pages: list[dict], out_path: Path):
    """Convert extracted structure into a properly formatted DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    doc = Document()

    # Narrow margins like a resume
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Remove default paragraph spacing
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(2)

    FONT_SIZES = {
        "large":  14,
        "medium": 11,
        "small":  10,
        "tiny":   9,
    }

    ALIGN_MAP = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }

    def add_runs(para, text: str, bold_parts: list, italic_parts: list,
                 font_size: int, default_bold=False):
        """Add runs to paragraph, applying bold/italic to matching phrases."""
        if not text:
            return

        # Build list of (start, end, bold, italic) spans
        # Simple approach: check each word/phrase
        remaining = text
        pos = 0

        # Tokenize by bold/italic markers
        # We'll do word-by-word matching
        words = re.split(r'(\s+)', text)
        current_run_text = ""
        current_bold = default_bold
        current_italic = False

        def flush_run(t, b, i):
            if not t:
                return
            run = para.add_run(t)
            run.bold = b
            run.italic = i
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"

        i = 0
        while i < len(words):
            word = words[i]
            word_stripped = word.strip()

            is_bold   = default_bold or any(bp in text[max(0,text.find(word_stripped)-2):text.find(word_stripped)+len(word_stripped)+2]
                                            for bp in bold_parts if bp and word_stripped and bp in word_stripped or word_stripped in bp)
            is_italic = any(ip in word_stripped or word_stripped in ip
                            for ip in italic_parts if ip)

            if is_bold != current_bold or is_italic != current_italic:
                flush_run(current_run_text, current_bold, current_italic)
                current_run_text = word
                current_bold = is_bold
                current_italic = is_italic
            else:
                current_run_text += word

            i += 1

        flush_run(current_run_text, current_bold, current_italic)

        # If no runs were added, add the full text as one run
        if not para.runs:
            run = para.add_run(text)
            run.bold = default_bold
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"

    def set_indent(para, level: int):
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.25 * level)

    for page_data in all_pages:
        sections = page_data.get("sections", [])

        for item in sections:
            stype       = item.get("type", "text")
            content     = item.get("content", "").strip()
            bold_parts  = item.get("bold_parts", []) or []
            italic_parts= item.get("italic_parts", []) or []
            font_size   = FONT_SIZES.get(item.get("font_size", "small"), 10)
            alignment   = item.get("alignment", "left")
            indent      = item.get("indent_level", 0) or 0
            is_bullet   = item.get("is_bullet", False)
            bullet_char = item.get("bullet_char", "–")
            columns     = item.get("columns", None)

            if not content and not columns:
                doc.add_paragraph("")
                continue

            # ── Table row (two columns: name | date/location) ──────
            if stype == "table_row" and columns and len(columns) >= 2:
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                # Remove borders
                from docx.oxml.ns import qn as _qn
                tbl = table._tbl
                tblPr = tbl.find(_qn("w:tblPr"))
                if tblPr is None:
                    tblPr = OxmlElement("w:tblPr")
                    tbl.insert(0, tblPr)
                tblBorders = OxmlElement("w:tblBorders")
                for border_name in ["top","left","bottom","right","insideH","insideV"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(_qn("w:val"), "none")
                    tblBorders.append(border)
                tblPr.append(tblBorders)

                # Set column widths
                row = table.rows[0]
                row.cells[0].width = Inches(4.5)
                row.cells[1].width = Inches(2.5)

                # Left cell
                lp = row.cells[0].paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                lr = lp.add_run(columns[0])
                lr.bold = True
                lr.font.size = Pt(font_size)
                lr.font.name = "Calibri"

                # Right cell
                rp = row.cells[1].paragraphs[0]
                rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                rr = rp.add_run(columns[1])
                rr.font.size = Pt(font_size)
                rr.font.name = "Calibri"
                continue

            # ── Section title (all caps / large) ───────────────────
            if stype == "section_title":
                p = doc.add_paragraph()
                p.alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
                # Add a top border (horizontal rule above section)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                top = OxmlElement("w:top")
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "6")
                top.set(qn("w:space"), "1")
                top.set(qn("w:color"), "auto")
                pBdr.append(top)
                pPr.append(pBdr)
                run = p.add_run(content.upper())
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Calibri"
                continue

            # ── Header (name / top of resume) ──────────────────────
            if stype == "header":
                p = doc.add_paragraph()
                p.alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(max(font_size, 16))
                run.font.name = "Calibri"
                continue

            # ── Bullet point ────────────────────────────────────────
            if is_bullet:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent  = Inches(0.25 + 0.25 * indent)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                # Add bullet char + space manually for exact match
                full_text = f"{bullet_char} {content}" if not content.startswith(bullet_char) else content
                add_runs(p, full_text, bold_parts, italic_parts, font_size)
                continue

            # ── Regular paragraph ────────────────────────────────────
            p = doc.add_paragraph()
            p.alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            set_indent(p, indent)

            is_default_bold = stype == "entry" or font_size >= 12
            add_runs(p, content, bold_parts, italic_parts, font_size, is_default_bold)

    doc.save(str(out_path))
    print(f"\n✅ Saved: {out_path.name} ({out_path.stat().st_size:,} bytes)")


# ── Main ─────────────────────────────────────────────────────────────

def convert(pdf_path_str: str):
    pdf = Path(pdf_path_str)
    if not pdf.exists():
        print(f"❌ File not found: {pdf_path_str}")
        sys.exit(1)

    out = pdf.parent / (pdf.stem + "_gpt4o.docx")
    t0  = time.time()

    print(f"\n{'='*50}")
    print(f"📄 PDF  : {pdf.name}")
    print(f"📝 Output: {out.name}")
    print(f"{'='*50}\n")

    # 1. Render pages
    print("Step 1: Rendering PDF pages…")
    images = pdf_to_images(pdf)
    print(f"  → {len(images)} page(s)\n")

    # 2. Extract structure via GPT-4o
    print("Step 2: Extracting structure via GPT-4o Vision…")
    all_pages = []
    for i, img in enumerate(images, 1):
        page_data = extract_structure_gpt4o(img, i)
        all_pages.append(page_data)
        n = len(page_data.get("sections", []))
        print(f"  → Page {i}: {n} sections extracted\n")

    # 3. Build DOCX
    print("Step 3: Building DOCX…")
    build_docx(all_pages, out)

    print(f"\n🎉 Total time: {time.time()-t0:.1f}s")
    print(f"📂 Open: {out.resolve()}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python test_gpt4o_pdf.py <path_to_resume.pdf>")
        sys.exit(1)
    convert(sys.argv[1])