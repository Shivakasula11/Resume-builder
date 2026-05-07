"""
Test script — Docling PDF→DOCX conversion
==========================================
Install:
    pip install docling python-docx

Run:
    python test_docling.py yourresume.pdf

Output:
    yourresume_docling.docx  — open in Word to check quality
"""

import sys
import time
from pathlib import Path


def convert(pdf_path: str):
    pdf = Path(pdf_path)
    if not pdf.exists():
        print(f"❌ File not found: {pdf_path}")
        sys.exit(1)

    print(f"📄 Converting: {pdf.name}")
    print("⏳ Loading Docling (first run downloads ML models ~500MB, be patient)…")

    t0 = time.time()

    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.base_models import InputFormat
    from docling.backend.docling_parse_v2_backend import DoclingParseV2DocumentBackend

    # Pipeline options — enable table structure + layout analysis
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False          # set True if PDF is scanned
    pipeline_options.do_table_structure = True
    pipeline_options.table_structure_options.do_cell_matching = True

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(
                pipeline_options=pipeline_options,
                backend=DoclingParseV2DocumentBackend,
            )
        }
    )

    print("🔄 Converting…")
    result = converter.convert(str(pdf))
    doc = result.document

    t1 = time.time()
    print(f"✓ Docling parsed in {t1-t0:.1f}s")

    # Export to DOCX
    out_path = pdf.parent / (pdf.stem + "_docling.docx")

    # Try native docx export first
    try:
        from docling.document_converter import WordFormatOption
        doc.save_as_docx(str(out_path))
        print(f"✓ Saved via native DOCX export: {out_path.name}")
    except Exception as e:
        print(f"⚠ Native DOCX export failed ({e}), falling back to markdown→docx…")
        _fallback_via_markdown(doc, out_path)

    size = out_path.stat().st_size
    print(f"\n✅ Done in {time.time()-t0:.1f}s — {out_path.name} ({size:,} bytes)")
    print(f"📂 Open: {out_path.resolve()}")


def _fallback_via_markdown(doc, out_path: Path):
    """Convert Docling document to DOCX via markdown as fallback."""
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md = doc.export_to_markdown()
    docx = DocxDoc()

    # Set narrow margins
    from docx.shared import Inches
    for section in docx.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    for line in md.split("\n"):
        stripped = line.strip()
        if not stripped:
            docx.add_paragraph("")
            continue

        if stripped.startswith("### "):
            p = docx.add_paragraph(stripped[4:])
            p.style = "Heading 3"
        elif stripped.startswith("## "):
            p = docx.add_paragraph(stripped[3:])
            p.style = "Heading 2"
        elif stripped.startswith("# "):
            p = docx.add_paragraph(stripped[2:])
            p.style = "Heading 1"
        elif stripped.startswith(("- ", "* ", "• ")):
            docx.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = docx.add_paragraph()
            run = p.add_run(stripped.strip("**"))
            run.bold = True
        else:
            # Handle inline bold **text** within lines
            p = docx.add_paragraph()
            parts = stripped.split("**")
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)

    docx.save(str(out_path))
    print(f"✓ Saved via markdown fallback: {out_path.name}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python test_docling.py <path_to_resume.pdf>")
        sys.exit(1)
    convert(sys.argv[1])