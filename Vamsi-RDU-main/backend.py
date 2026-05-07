"""
Resume Editor Backend v7.6
===========================
PDF  upload → Adobe PDF Services (best quality) → current.docx → edit → download as DOCX/PDF
DOCX upload → current.docx → edit directly → download as DOCX/PDF

Conversion priority:
  PDF→DOCX : 1) Adobe PDF Services  2) MS Word COM  3) pdf2docx
  DOCX→PDF : 1) docx2pdf            2) LibreOffice

Key principle — SURGICAL editing:
  Text:       only w:t nodes changed, w:rPr (font/size/color/bold/spacing) untouched
  Formatting: only properties the user explicitly set are written; nothing else changed

v7.6 changes:
  - add-page now returns page_break_para_index so frontend can target the new page
  - new /write-page-block endpoint: writes multi-line plain text to a blank page
    (one paragraph per newline, inserted after the page-break paragraph)
"""

import os
from dotenv import load_dotenv
load_dotenv()

import re
import uuid
import shutil
import subprocess
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import List, Optional

import requests
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from lxml import etree

from docx import Document
from docx.oxml.ns import qn

# ── optional converters ──────────────────────────────────────────────
try:
    from pdf2docx import Converter as _PDF2DOCX
    PDF2DOCX_OK = True
except ImportError:
    PDF2DOCX_OK = False

try:
    import comtypes.client
    WORD_COM_OK = True
except ImportError:
    WORD_COM_OK = False

try:
    from docx2pdf import convert as _docx2pdf
    DOCX2PDF_OK = True
except ImportError:
    DOCX2PDF_OK = False

ILOVEPDF_KEY  = os.environ.get("ILOVEPDF_PUBLIC_KEY", "").strip()
ILOVEPDF_OK   = False
ADOBE_CID     = os.environ.get("ADOBE_CLIENT_ID", "").strip()
ADOBE_CSECRET = os.environ.get("ADOBE_CLIENT_SECRET", "").strip()
ADOBE_OK      = bool(ADOBE_CID and ADOBE_CSECRET)

# ── app ──────────────────────────────────────────────────────────────
app = FastAPI(title="Resume Editor API v7.6")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

try:
    from ai_routes import ai_router
    app.include_router(ai_router, prefix="/ai", tags=["AI"])
    AI_SUPPORT = True
except ImportError:
    AI_SUPPORT = False

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

print(f"""
[Resume Editor v7.6]
  Adobe PDF Services  : {'✓ client=' + ADOBE_CID[:12] + '…' if ADOBE_OK else '✗  set ADOBE_CLIENT_ID + ADOBE_CLIENT_SECRET'}
  Word COM (PDF→DOCX) : {'✓' if WORD_COM_OK else '✗'}
  pdf2docx (PDF→DOCX) : {'✓' if PDF2DOCX_OK else '✗'}
  docx2pdf (DOCX→PDF) : {'✓' if DOCX2PDF_OK else '✗'}
  AI module           : {'✓' if AI_SUPPORT else '✗'}
""")


# ════════════════════════════════════════════════════════════════════
#  PDF → DOCX
# ════════════════════════════════════════════════════════════════════

_ICON_REPLACEMENTS = {
    "II": "in", "Im": "✉", "O": "gh", "f": "fb", "t": "tw",
    "B": "be", "D": "dr", "": "in", "": "✉", "": "gh",
}

_ICON_FONTS  = {"arialmt", "arial", "symbol", "wingdings", "webdings"}
_HEAVY_FONTS = {"georgia", "georgia-bold", "georgiabold"}
_LIGHT_REPLACEMENT = "Calibri"


def _force_black_all_runs(docx_path: Path, para_index: int):
    try:
        doc = Document(str(docx_path))
        paras = doc.paragraphs
        if para_index < 0 or para_index >= len(paras):
            return
        p_elem = paras[para_index]._element
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            ppr_rpr = pPr.find(qn("w:rPr"))
            if ppr_rpr is not None:
                pPr.remove(ppr_rpr)
        for r in p_elem.iter(qn("w:r")):
            if r.getparent() is not None and r.getparent().tag == qn("w:hyperlink"):
                continue
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = r.makeelement(qn("w:rPr"), {})
                r.insert(0, rpr)
            col = rpr.find(qn("w:color"))
            if col is None:
                col = rpr.makeelement(qn("w:color"), {})
                rpr.append(col)
            col.set(qn("w:val"), "000000")
            for attr in list(col.attrib.keys()):
                if attr != qn("w:val"): del col.attrib[attr]
            rs = rpr.find(qn("w:rStyle"))
            if rs is not None: rpr.remove(rs)
        doc.save(str(docx_path))
    except Exception as e:
        print(f"[ForceBlack] {e}")


def _patch_styles_gray(docx_path: Path):
    import zipfile, re, io
    try:
        buf = io.BytesIO()
        with zipfile.ZipFile(str(docx_path), 'r') as zin:
            with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == 'word/styles.xml':
                        xml = data.decode('utf-8')
                        xml = re.sub(r'<w:color\s+[^>]*w:themeColor[^/]*/?>','', xml)
                        xml = re.sub(
                            r'<w:color\s+w:val="(?:[Aa]6[Aa]6[Aa]6|7[Ff]7[Ff]7[Ff]|[Bb][Ff][Bb][Ff][Bb][Ff]|808080|595959|262626|[Dd]9[Dd]9[Dd]9|[Cc]0[Cc]0[Cc]0|[Aa]9[Aa]9[Aa]9)"\s*/?>',
                            '', xml)
                        data = xml.encode('utf-8')
                    zout.writestr(item, data)
        buf.seek(0)
        with open(str(docx_path), 'wb') as f:
            f.write(buf.read())
        print(f"[PatchStyles] ✓ gray colors removed from styles.xml")
    except Exception as e:
        print(f"[PatchStyles] warning: {e}")


def _fix_adobe_bold(docx_path: Path):
    doc = Document(str(docx_path))
    font_fixes = 0; icon_fixes = 0

    def _get_font(rpr):
        if rpr is None: return None
        fonts_elem = rpr.find(qn("w:rFonts"))
        if fonts_elem is None: return None
        return (fonts_elem.get(qn("w:ascii")) or
                fonts_elem.get(qn("w:hAnsi")) or
                fonts_elem.get(qn("w:cs")))

    def _set_font(rpr, font_name):
        fonts_elem = rpr.find(qn("w:rFonts"))
        if fonts_elem is None:
            fonts_elem = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, fonts_elem)
        for attr in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs")):
            fonts_elem.set(attr, font_name)

    def _is_bold(rpr):
        if rpr is None: return False
        b = rpr.find(qn("w:b"))
        return b is not None and b.get(qn("w:val"), "true") not in ("0", "false")

    def _process_runs(r_elems):
        nonlocal font_fixes, icon_fixes
        for r_elem in r_elems:
            text = "".join(t.text or "" for t in r_elem.findall(qn("w:t"))).strip()
            if not text: continue
            rpr = r_elem.find(qn("w:rPr"))
            font = (_get_font(rpr) or "").lower()
            bold = _is_bold(rpr)
            if font in _HEAVY_FONTS and not bold and rpr is not None:
                _set_font(rpr, _LIGHT_REPLACEMENT); font_fixes += 1
            if font in _ICON_FONTS and len(text) <= 3:
                replacement = _ICON_REPLACEMENTS.get(text)
                if replacement:
                    for t_elem in r_elem.findall(qn("w:t")):
                        t_elem.text = replacement
                        t_elem.set(qn("xml:space"), "preserve")
                    if rpr is None:
                        rpr = r_elem.makeelement(qn("w:rPr"), {})
                        r_elem.insert(0, rpr)
                    _set_font(rpr, _LIGHT_REPLACEMENT); icon_fixes += 1

    for para in doc.paragraphs:
        _process_runs(para._element.findall(qn("w:r")))
        for hl in para._element.findall(qn("w:hyperlink")):
            _process_runs(hl.findall(qn("w:r")))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_runs(para._element.findall(qn("w:r")))

    if font_fixes or icon_fixes:
        doc.save(str(docx_path))
        print(f"[Adobe] ✓ post-process: {font_fixes} font-weight fixes, {icon_fixes} icon fixes")


def _convert_adobe(pdf_path: Path, docx_path: Path) -> bool:
    try:
        from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
        from adobe.pdfservices.operation.pdf_services import PDFServices
        from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
        from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
        from adobe.pdfservices.operation.io.stream_asset import StreamAsset
        from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
        from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
        from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
        from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult

        import socket as _sock
        try:
            _s = _sock.socket(_sock.AF_INET, _sock.SOCK_STREAM)
            _s.settimeout(2)
            _s.connect(("pdf-services-ue1.adobe.io", 443))
            _s.close()
        except Exception:
            raise Exception("Adobe unreachable — skipping to next converter")

        credentials = ServicePrincipalCredentials(client_id=ADOBE_CID, client_secret=ADOBE_CSECRET)
        pdf_services = PDFServices(credentials=credentials)

        with open(pdf_path, "rb") as f:
            input_asset = pdf_services.upload(input_stream=f, mime_type=PDFServicesMediaType.PDF)

        params   = ExportPDFParams(target_format=ExportPDFTargetFormat.DOCX)
        job      = ExportPDFJob(input_asset=input_asset, export_pdf_params=params)
        location = pdf_services.submit(job)
        response = pdf_services.get_job_result(location, ExportPDFResult)

        result_asset: CloudAsset = response.get_result().get_asset()
        stream_asset: StreamAsset = pdf_services.get_content(result_asset)

        with open(docx_path, "wb") as f:
            data = stream_asset.get_input_stream()
            f.write(data if isinstance(data, bytes) else data.read())

        _fix_adobe_bold(docx_path)
        print(f"[Adobe] ✓  {docx_path.stat().st_size:,} bytes")
        return True

    except Exception as e:
        import traceback
        print(f"[Adobe] exception: {e}\n{traceback.format_exc()}")
        return False


def _convert_word_com(pdf_path: Path, docx_path: Path) -> bool:
    try:
        try:
            import winreg
            for ver in ["16.0","15.0","14.0"]:
                try:
                    with winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                          rf"Software\Microsoft\Office\{ver}\Word\Options",
                          0, winreg.KEY_SET_VALUE) as k:
                        winreg.SetValueEx(k,"DisableConvertPdfWarning",0,winreg.REG_DWORD,1)
                except FileNotFoundError:
                    pass
        except Exception:
            pass
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False; word.DisplayAlerts = 0
        word.AutomationSecurity = 3
        word.Options.ConfirmConversions = False
        doc = word.Documents.Open(
            FileName=str(pdf_path.resolve()),
            ConfirmConversions=False, ReadOnly=True,
            AddToRecentFiles=False, NoEncodingDialog=True)
        doc.SaveAs2(str(docx_path.resolve()), FileFormat=16)
        doc.Close(False); word.Quit()
        if docx_path.exists():
            print("[Word COM] ✓"); return True
    except Exception as e:
        print(f"[Word COM] failed: {e}")
        try: word.Quit()
        except: pass
    return False


def _convert_pdf2docx(pdf_path: Path, docx_path: Path) -> bool:
    try:
        cv = _PDF2DOCX(str(pdf_path))
        cv.convert(str(docx_path)); cv.close()
        if docx_path.exists():
            print("[pdf2docx] ✓"); return True
    except Exception as e:
        print(f"[pdf2docx] failed: {e}")
    return False


def pdf_to_docx(pdf_path: Path, docx_path: Path) -> tuple:
    if ADOBE_OK:
        if _convert_adobe(pdf_path, docx_path): return True, "adobe"
        print("[PDF→DOCX] Adobe failed, trying next…")
    if WORD_COM_OK:
        if _convert_word_com(pdf_path, docx_path): return True, "word_com"
        print("[PDF→DOCX] Word COM failed, trying next…")
    if PDF2DOCX_OK:
        if _convert_pdf2docx(pdf_path, docx_path): return True, "pdf2docx"
    return False, "none"


# ════════════════════════════════════════════════════════════════════
#  DOCX → PDF
# ════════════════════════════════════════════════════════════════════

def docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    if pdf_path.exists(): pdf_path.unlink()

    if DOCX2PDF_OK:
        try:
            _docx2pdf(str(docx_path.resolve()), str(pdf_path.resolve()))
            if pdf_path.exists():
                print("[DOCX→PDF] docx2pdf ✓"); return True
        except Exception as e:
            print(f"[DOCX→PDF] docx2pdf: {e}")

    for cmd in ["libreoffice","soffice"]:
        try:
            subprocess.run([cmd,"--headless","--convert-to","pdf",
                           "--outdir", str(docx_path.parent), str(docx_path)],
                          capture_output=True, timeout=90)
            gen = docx_path.parent / (docx_path.stem + ".pdf")
            if gen.exists():
                gen.rename(pdf_path); print(f"[DOCX→PDF] {cmd} ✓"); return True
        except FileNotFoundError: continue
        except Exception as e: print(f"[DOCX→PDF] {cmd}: {e}")

    return False


# ════════════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ════════════════════════════════════════════════════════════════════

def flatten_hyperlinks(doc_path: str, out_path=None):
    doc = Document(doc_path); changed = False
    def _rt(r): return "".join(t.text or "" for t in r.findall(qn("w:t"))).strip()
    def _ri(r): return r.find(qn("w:drawing")) is not None
    def _mkhl(tmpl, runs):
        hl = deepcopy(tmpl)
        for ch in list(hl): hl.remove(ch)
        for r in runs: hl.append(r)
        return hl
    def _prom(hl):
        out=[]; runs=[]; pi=False
        for ch in hl:
            if ch.tag==qn("w:hyperlink"):
                if runs: out.append(_mkhl(hl,runs)); runs=[]
                pi=False; out.extend(_prom(ch))
            elif ch.tag==qn("w:r"):
                if pi or _ri(ch):
                    if runs: out.append(_mkhl(hl,runs)); runs=[]
                    out.append(deepcopy(ch))
                    if _ri(ch): pi=True
                elif _rt(ch): runs.append(deepcopy(ch))
                else:
                    if runs: runs.append(deepcopy(ch))
                    else: out.append(deepcopy(ch))
        if runs: out.append(_mkhl(hl,runs))
        return out
    for para in doc.paragraphs:
        p=para._element
        for hl in list(p):
            if hl.tag!=qn("w:hyperlink") or not hl.findall(qn("w:hyperlink")): continue
            flat=_prom(hl); ins=hl
            for el in flat: ins.addnext(el); ins=el
            p.remove(hl); changed=True
    if changed: doc.save(out_path or doc_path)


def _get_runs(para):
    results=[]
    def _p(r_elem, is_hl=False, ro=None):
        text=""; has_tab=has_br=has_draw=False
        for s in r_elem:
            if s.tag==qn("w:t"):         text+=s.text or ""
            elif s.tag==qn("w:tab"):     has_tab=True
            elif s.tag==qn("w:br"):      has_br=True
            elif s.tag==qn("w:drawing"): has_draw=True
        bold=False
        rpr=r_elem.find(qn("w:rPr"))
        if rpr is not None:
            b=rpr.find(qn("w:b"))
            if b is not None: bold=b.get(qn("w:val")) not in ("0","false")
        results.append(dict(run=ro,text=text,is_hl=is_hl,elem=r_elem,
                            bold=bold,has_tab=has_tab,has_br=has_br,has_draw=has_draw))
    for ch in para._element:
        if ch.tag==qn("w:r"):
            ro=next((r for r in para.runs if r._element is ch),None); _p(ch,False,ro)
        elif ch.tag==qn("w:hyperlink"):
            for hr in ch.findall(qn("w:r")): _p(hr,True)
        elif ch.tag==qn("w:sdt"):
            sc=ch.find(qn("w:sdtContent"))
            if sc:
                for c in sc:
                    if c.tag==qn("w:r"): _p(c)
                    elif c.tag==qn("w:hyperlink"):
                        for hr in c.findall(qn("w:r")): _p(hr,True)
    return results


def _para_text(para):
    return "".join(r["text"] for r in _get_runs(para))


def _run_fmt(ri):
    parts=["bold"] if ri["bold"] else []
    if ri["is_hl"]: parts.append("hyperlink")
    r=ri.get("run")
    if r:
        try:
            if r.italic: parts.append("italic")
        except: pass
        try:
            if r.font.size: parts.append(f"{r.font.size.pt:.0f}pt")
        except: pass
        try:
            if r.font.name: parts.append(r.font.name)
        except: pass
    return ", ".join(parts) or "regular"


def _can_merge(a,b):
    if any(a.get(k) or b.get(k) for k in ("has_tab","has_br","has_draw")): return False
    if a["bold"]!=b["bold"] or a["is_hl"]!=b["is_hl"]: return False
    ra,rb=a.get("run"),b.get("run")
    if ra and rb:
        try:
            if ra.italic!=rb.italic: return False
            if ra.font.size!=rb.font.size: return False
            if ra.font.name!=rb.font.name: return False
        except: pass
    return True


def _run_groups(para):
    all_r = _get_runs(para)
    if not all_r: return [], all_r

    segments = []; cur_text = ""; cur_indices = []
    for i, ri in enumerate(all_r):
        if ri["has_br"] or ri["has_draw"]:
            if cur_text.strip():
                segments.append((cur_text, cur_indices[:]))
            cur_text = ""; cur_indices = []; continue
        if ri.get("has_tab") and not ri["text"].strip():
            if cur_text.strip():
                cur_text += "  |  "
            continue
        cur_text += ri["text"]
        if ri["text"]: cur_indices.append(i)

    if cur_text.strip():
        segments.append((cur_text, cur_indices[:]))

    if not segments: return [], all_r

    full_text = " ".join(s[0] for s in segments).strip()
    all_indices = [i for seg in segments for i in seg[1]]
    if not full_text: return [], all_r

    first = next((all_r[i] for i in all_indices if all_r[i]["text"].strip()), None)
    is_bold = first["bold"] if first else False
    has_hl  = any(all_r[i]["is_hl"] for i in all_indices)
    fmt     = _run_fmt(first) if first else "regular"

    return [dict(text=full_text, indices=all_indices,
                 fmt=fmt, is_bold=is_bold, has_hl=has_hl)], all_r


# ─── section / field detection ───────────────────────────────────────

_SECS=[
    dict(key="summary",        label="Professional Summary",color="#7c3aed",
         pat=r"^(professional\s*)?summary|profile|objective|about\s*me"),
    dict(key="skills",         label="Skills",              color="#059669",
         pat=r"^(technical\s*)?skills|competenc|technolog"),
    dict(key="experience",     label="Work Experience",      color="#2563eb",
         pat=r"^(work\s*)?experience|employment|work\s*history"),
    dict(key="projects",       label="Projects",            color="#dc2626",
         pat=r"^projects?"),
    dict(key="education",      label="Education",           color="#d97706",
         pat=r"^education|academic|qualifications?"),
    dict(key="certifications", label="Certifications",      color="#0891b2",
         pat=r"^certifications?|certificates?"),
    dict(key="contact",        label="Contact",             color="#6366f1",
         pat=r"^contact"),
]

def _normalize_sec_text(text):
    """Remove Adobe letter-spacing artifact: 'S K I L L S' → 'SKILLS'."""
    t = text.strip()
    # Pattern: single uppercase letters separated by spaces (3+ chars total)
    if re.match(r'^([A-Za-z] ){2,}[A-Za-z]$', t):
        return t.replace(' ', '')
    return t

def _detect_sec(text, style_name=None):
    t = _normalize_sec_text(text.strip())
    if len(t) > 50: return None
    is_h=style_name and ("heading" in style_name.lower() or "title" in style_name.lower())
    for s in _SECS:
        if re.match(s["pat"],t,re.IGNORECASE):
            if is_h: return s
            if re.sub(r"[:\s]+$","",t).lower()==t.strip().lower(): return s
            return None
    return None

def _classify(text,sec_key,fid):
    if re.search(r"[\w.-]+@[\w.-]+\.\w+",text): return "email"
    if re.search(r"^\+?\d[\d\s\-()]{8,}",text): return "phone"
    if re.search(r"linkedin|github",text,re.IGNORECASE): return "link"
    if sec_key=="header" and fid<=2: return "name" if fid==0 else "title"
    return "text"


# ─── field extraction ────────────────────────────────────────────────





# ─── LLM-powered section classification ─────────────────────────────────────
#
# Single GPT-4o-mini call per resume upload.
# Reads every extracted field and assigns the correct section key semantically.
# Falls back to the regex pipeline if the API is unavailable or returns bad data.


def _classify_sections_llm(fields):
    """
    Primary: GPT-4o-mini semantic classification — order-independent, language-aware,
    handles all layout types (two-column, table-based, non-English, unusual headings).

    Fallback: three-stage regex pipeline (content patterns → project zones → proximity).
    Only activates if the API call fails or returns a malformed response.
    """
    import os, json as _json
    key = os.environ.get("OPENAI_API_KEY", "").strip()
    if key:
        result = _llm_classify(fields, key)
        if result is not None:
            return result
    # API unavailable or failed → regex fallback
    return _regex_classify(fields)


def _llm_classify(fields, api_key):
    """
    Call GPT-4o-mini with response_format=json_object to guarantee valid JSON.
    Returns None (triggers regex fallback) if quality is too low.
    """
    import json as _json
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)

        classifiable = [
            (i, f) for i, f in enumerate(fields)
            if f.get("source") == "body"
            and not f.get("isTextbox")
            and f.get("type") != "image"
            and f.get("text", "").strip()
        ]
        if not classifiable:
            return fields

        n = len(classifiable)
        valid_keys = {"header","summary","skills","experience","education",
                      "projects","certifications","languages","awards","other"}

        def clean(text, maxlen=160):
            t = text[:maxlen].replace('"',' ').replace("'",' ').replace('\\',' ')
            return re.sub(r'[\x00-\x1f]', ' ', t).strip()

        # Build field lines with preceding-header context so LLM understands
        # which section each field sits under in the document structure
        headers_in_order = sorted(
            [(f["paraIndex"], f["section"]["key"])
             for f in fields if f.get("isHeader") and f.get("source")=="body"
             and f.get("paraIndex") is not None],
            key=lambda x: x[0]
        )
        def get_context(pidx):
            prec = [h for h in headers_in_order if h[0] <= pidx]
            return prec[-1][1] if prec else "start"

        lines_out = []
        for li, (_, f) in enumerate(classifiable):
            tag = "HEADING" if f.get("isHeader") else f"under:{get_context(f.get('paraIndex',0))}"
            lines_out.append(f"{li+1}. [{tag}] {clean(f['text'])}")
        field_lines = "\n".join(lines_out)

        valid_section_keys = ["header","summary","skills","experience","education",
                               "projects","certifications","languages","awards","other"]
        system_msg = (
            "You are a resume section classifier. "
            "I will give you numbered resume fields. "
            "For EACH field, output exactly ONE classification label. "
            "The labels are: header, summary, skills, experience, education, "
            "projects, certifications, languages, awards, other. "
            "IMPORTANT: Output ONLY these label words — never copy the field text into the output. "
            "Return a JSON object: {\"sections\": [\"label1\", \"label2\", ...]} "
            "where every element is one of the 10 label words above, nothing else. "
            "Example output for 4 fields: {\"sections\": [\"header\", \"summary\", \"skills\", \"experience\"]}. "
            "Classification rules: "
            "Candidate name / email / phone / LinkedIn / GitHub profile URL = header. "
            "Professional summary or objective paragraph = summary. "
            "Technical/soft/ML skills lines = skills. "
            "Job titles / company names / work dates / job description bullets = experience. "
            "Personal or academic project titles, project GitHub links, project description bullets = projects. "
            "Certificates / courses / virtual internships / workshops = certifications. "
            "Degrees / school names / CGPA / grades = education. "
            "Section heading lines (SKILLS, EXPERIENCE, PROJECTS etc.) → classify to their own section."
        )

        user_msg = (
            f"Classify each of the {n} resume fields below. "
            f"Return JSON: {{\"sections\": [label1, label2, ...]}} with exactly {n} label words.\n"
            f"Each label MUST be one of: {', '.join(valid_section_keys)}.\n"
            "Do NOT include any field text in your response — only the label words.\n\n"
            + field_lines
        )

        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": user_msg},
            ],
            temperature=0,
            max_tokens=max(400, n * 15),
            response_format={"type": "json_object"},
        )

        raw = resp.choices[0].message.content.strip()

        # Parse JSON — response_format guarantees valid JSON on openai>=1.0
        # For older library versions, apply repair before parsing
        try:
            data = _json.loads(raw)
        except _json.JSONDecodeError:
            # Strip markdown fences
            raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.MULTILINE).strip()
            # Extract the first JSON object or array
            m = re.search(r'(\{.*\}|\[.*\])', raw, re.DOTALL)
            if m:
                raw = m.group(1)
            # Remove control characters and fix common issues
            raw = re.sub(r'[\x00-\x08\x0b-\x1f\x7f]', '', raw)
            raw = re.sub(r',\s*([\]\}])', r'\1', raw)  # trailing commas
            try:
                data = _json.loads(raw)
            except _json.JSONDecodeError:
                # Extract section keys directly via regex as last resort
                keys = re.findall(r'"(header|summary|skills|experience|education|'
                                  r'projects|certifications|languages|awards|other)"', raw)
                if len(keys) >= n * 0.8:
                    data = {"sections": keys}
                else:
                    print(f"[LLM] JSON unrecoverable — using regex fallback")
                    return None

        if isinstance(data, dict):
            assigned = (data.get("sections")
                        or data.get("classifications")
                        or next(iter(data.values()), None))
        else:
            assigned = data

        if not isinstance(assigned, list):
            print("[LLM] unexpected response shape — using regex fallback")
            return None

        # Too short → don't pad, let regex handle it cleanly
        if len(assigned) < n:
            print(f"[LLM] short by {n - len(assigned)} — using regex fallback")
            return None

        assigned = assigned[:n]  # trim if too long

        # Sanity check: if values look like field text (avg length > 10 chars),
        # LLM returned field contents instead of labels — fall back immediately
        avg_len = sum(len(str(k)) for k in assigned) / max(len(assigned), 1)
        if avg_len > 15:
            print(f"[LLM] returned field text instead of labels (avg_len={avg_len:.1f}) — regex fallback")
            return None

        # Quality guard: if >40% are "other", classification failed
        other_count = sum(1 for k in assigned
                          if str(k).strip().lower() not in valid_keys
                          or str(k).strip().lower() == "other")
        if other_count / n > 0.40:
            print(f"[LLM] too many other ({other_count}/{n}) — using regex fallback")
            return None

        _sec_map = {s["key"]: s for s in _SECS}
        _sec_map.update({
            "header":  {"key":"header",  "label":"Header",  "color":"#c026d3"},
            "other":   {"key":"other",   "label":"Other",   "color":"#6b7280"},
            "contact": {"key":"contact", "label":"Contact", "color":"#6366f1"},
        })

        reassigned = 0
        for (field_idx, field), sec_key in zip(classifiable, assigned):
            sec_key = str(sec_key).strip().lower()
            if sec_key not in valid_keys:
                sec_key = "other"
            new_sec = _sec_map.get(sec_key, _sec_map["other"])
            if new_sec["key"] != field.get("section", {}).get("key"):
                fields[field_idx] = dict(field, section=dict(new_sec))
                reassigned += 1

             # ── Post-LLM correction pass ──────────────────────────────────
        corrected = 0
        github_re_post = re.compile(r"github\.com|https?://github", re.IGNORECASE)
        title_re_post  = re.compile(r".{3,}\s[\u2013\u2014\-]\s.{3,}", re.UNICODE)
        _proj_sec_post = _SEC_DICT.get("projects", {"key":"projects","label":"Projects","color":"#dc2626"})

        headers_sorted = sorted(
            [(f["paraIndex"], f["section"]["key"])
             for f in fields if f.get("isHeader") and f.get("source") == "body"
             and f.get("paraIndex") is not None],
            key=lambda x: x[0]
        )
        def _prec_key(pidx):
            prec = [h for h in headers_sorted if h[0] <= pidx]
            return prec[-1][1] if prec else None

        # Stage A: strong content patterns override wrong LLM guesses
        touched_post = set()
        for i, field in enumerate(fields):
            if field.get("isHeader") or field.get("source") != "body":
                continue
            text = field.get("text", "").strip()
            if not text:
                continue
            current = field.get("section", {}).get("key", "")
            pidx = field.get("paraIndex")
            prec = _prec_key(pidx) if pidx is not None else None
            for sec_key, pat in _STRONG_PATTERNS:
                if pat.search(text):
                    if sec_key == "projects" and prec == "experience":
                        continue
                    if sec_key != current:
                        fields[i] = dict(field, section=dict(_SEC_DICT[sec_key]))
                        corrected += 1
                        touched_post.add(i)
                    break
            curr_now = fields[i].get("section", {}).get("key", "")
            if (curr_now == "projects" and prec == "experience"
                    and re.search(
                        r"\b(into|via|for\s+[A-Z]|platform|imentora|imeetpro|lanciere|"
                        r"the\s+system|the\s+platform|the\s+application)\b",
                        text, re.IGNORECASE)):
                fields[i] = dict(fields[i], section=dict(_SEC_DICT["experience"]))
                touched_post.add(i)
                corrected += 1
            if (fields[i].get("section", {}).get("key") == "education"
                    and re.search(r"\binternship\b", text, re.IGNORECASE)):
                fields[i] = dict(fields[i], section=dict(_SEC_DICT["certifications"]))
                touched_post.add(i)
                corrected += 1

        # Stage B: project-zone detection — catches projects when PROJECTS header
        # is in a table cell and LLM couldn't see it, classifying them as experience
        gh_paras_post = set()
        for f in fields:
            if f.get("source") == "body" and github_re_post.search(f.get("text", "")):
                p = f.get("paraIndex", -1)
                for dv in range(-2, 4):
                    gh_paras_post.add(p + dv)
        proj_title_pidxs_post = set()
        for f in fields:
            if f.get("source") != "body" or f.get("isHeader"):
                continue
            pidx_t = f.get("paraIndex", -1)
            text_t = f.get("text", "").strip()
            if title_re_post.search(text_t) and len(text_t) < 90 and pidx_t in gh_paras_post:
                proj_title_pidxs_post.add(pidx_t)
        if proj_title_pidxs_post:
            sorted_titles = sorted(proj_title_pidxs_post)
            end_pidx = min(
                (h[0] for h in headers_sorted if h[1] in ("education", "certifications")),
                default=max(proj_title_pidxs_post) + 50
            )
            zones = [(tp, sorted_titles[ti+1] if ti+1 < len(sorted_titles) else end_pidx)
                     for ti, tp in enumerate(sorted_titles)]
            for i, field in enumerate(fields):
                if i in touched_post or field.get("source") != "body" or field.get("isHeader"):
                    continue
                pidx_z = field.get("paraIndex")
                if pidx_z is None:
                    continue
                for zs, ze in zones:
                    if zs <= pidx_z < ze and field.get("section", {}).get("key") != "projects":
                        fields[i] = dict(field, section=dict(_proj_sec_post))
                        corrected += 1
                        break

        print(f"[LLM] \u2713 {n} fields classified, {reassigned} reassigned, {corrected} pattern-corrected")
        return fields

    except Exception as e:
        print(f"[LLM] failed: {e} — using regex fallback")
        return None
# ── Module-level constants used by both LLM and regex classifiers ────────────
_STRONG_PATTERNS = [
    ("skills", re.compile(
        r"^(technical|soft|core|key|hard|ml.?ai|programming|language|tool|framework|"
        r"technology|professional|interpersonal|computer|analytical)[\s\-:]*(skills?|competenc|"
        r"frameworks?|tools?|stack|knowledge|abilities?|proficienc)"
        r"|^skills?\s*:",
        re.IGNORECASE)),
    ("education", re.compile(
        r"\b(b\.?\s*tech|b\.?\s*e\.?|m\.?\s*tech|m\.?\s*s\.?|b\.?\s*sc|m\.?\s*sc|mba|phd|"
        r"bachelor|degree|cgpa|gpa|percentage|university|college|school|institute|"
        r"intermediate|secondary|hsc|ssc|10th|12th)\b",
        re.IGNORECASE)),
    ("certifications", re.compile(
        r"(virtual internship|virtual intern|data science.*intern|"
        r"ai.ml.*intern|machine learning.*intern|certified|certification|"
        r"certificate|coursera|udemy|nptel|edx|workshop|bootcamp)",
        re.IGNORECASE)),
    ("projects", re.compile(
        r"^(built|developed|designed|created|deployed|engineered|automated|architected|constructed)\b"
        r"(?!.{0,80}\b(into|via|for\s+[A-Z]|using.*platform|imentora|imeetpro|lanciere|the\s+platform))",
        re.IGNORECASE)),
]

_SEC_DICT = {s["key"]: s for s in _SECS}
_SEC_DICT.update({
    "header":  {"key":"header",  "label":"Header",  "color":"#c026d3"},
    "other":   {"key":"other",   "label":"Other",   "color":"#6b7280"},
    "contact": {"key":"contact", "label":"Contact", "color":"#6366f1"},
})


def _regex_classify(fields):
    """Three-stage regex fallback: content patterns → project zones → preceding-header."""
    import re as _re

    headers = [
        (f["paraIndex"], f["section"])
        for f in fields
        if f.get("isHeader") and f.get("source") == "body"
        and f.get("paraIndex") is not None
    ]

    def preceding_sec(pidx):
        prec = [h for h in headers if h[0] <= pidx]
        return prec[-1][1]["key"] if prec else None

    _proj_sec = _SEC_DICT.get("projects", {"key":"projects","label":"Projects","color":"#dc2626"})
    github_re = _re.compile(r"github\.com|https?://github", _re.IGNORECASE)
    title_re  = _re.compile(r".{3,}\s[\u2013\u2014\-]\s.{3,}", _re.UNICODE)

    # Stage A: strong content patterns
    touched_A = set()
    for i, field in enumerate(fields):
        if field.get("isHeader") or field.get("source") != "body":
            continue
        text = field.get("text", "").strip()
        if not text:
            continue
        current = field.get("section", {}).get("key", "")
        pidx = field.get("paraIndex")
        prec = preceding_sec(pidx) if pidx is not None else None
        matched = None
        for sec_key, pat in _STRONG_PATTERNS:
            if pat.search(text):
                if sec_key == "projects" and prec == "experience":
                    continue
                matched = sec_key
                break
        if matched == "education" and _re.search(r"\binternship\b", text, _re.IGNORECASE):
            matched = "certifications"
        if matched and matched != current:
            fields[i] = dict(field, section=dict(_SEC_DICT[matched]))
            touched_A.add(i)

    # Stage B: project zones via title+GitHub
    gh_paras = set()
    for f in fields:
        if f.get("source") == "body" and github_re.search(f.get("text", "")):
            p = f.get("paraIndex", -1)
            for d in range(-2, 4):
                gh_paras.add(p + d)
    proj_title_pidxs = set()
    for f in fields:
        if f.get("source") != "body" or f.get("isHeader"):
            continue
        pidx = f.get("paraIndex", -1)
        text = f.get("text", "").strip()
        if title_re.search(text) and len(text) < 90 and pidx in gh_paras:
            proj_title_pidxs.add(pidx)
    if proj_title_pidxs:
        sorted_titles = sorted(proj_title_pidxs)
        end_pidx = min(
            (h[0] for h in headers if h[1]["key"] in ("education", "certifications")),
            default=max(proj_title_pidxs) + 50
        )
        zones = [(tp, sorted_titles[ti+1] if ti+1 < len(sorted_titles) else end_pidx)
                 for ti, tp in enumerate(sorted_titles)]
        for i, field in enumerate(fields):
            if field.get("source") != "body" or field.get("isHeader"):
                continue
            pidx = field.get("paraIndex")
            if pidx is None:
                continue
            for zs, ze in zones:
                if zs <= pidx < ze and field.get("section", {}).get("key") != "projects":
                    fields[i] = dict(field, section=dict(_proj_sec))
                    break

    # Stage C: preceding-header fallback
    for i, field in enumerate(fields):
        if field.get("isHeader") or field.get("source") != "body":
            continue
        pidx = field.get("paraIndex")
        if pidx is None:
            continue
        current = field.get("section", {}).get("key", "")
        if i in touched_A or current == "projects":
            continue
        prec = preceding_sec(pidx)
        if prec and prec != current:
            fields[i] = dict(field, section=dict(
                next(h[1] for h in reversed(headers) if h[0] <= pidx)
            ))

    return fields


def extract_fields(doc_path: str) -> list:
    doc=Document(doc_path); fields=[]; fid=0
    cur_sec=dict(key="header",label="Header",color="#c026d3")

    rel_map={r.rId:r.target_ref for r in doc.part.rels.values()
             if "hyperlink" in str(r.reltype).lower()}

    icon_links={}; seen_rids=set()
    for pidx,para in enumerate(doc.paragraphs):
        for drawing in para._element.iter(qn("w:drawing")):
            xs=etree.tostring(drawing).decode()
            rids=re.findall(r'hlinkClick[^>]*r:id="(rId\d+)"',xs)
            if not rids: continue
            rid=rids[0]
            if rid in seen_rids: continue
            seen_rids.add(rid); url=rel_map.get(rid,"")
            if not url: continue
            label=("LinkedIn" if "linkedin" in url.lower()
                   else "GitHub"  if "github"   in url.lower()
                   else "Mail"    if url.lower().startswith("mailto:")
                   else "Link")
            icon_links.setdefault(pidx,[]).append(dict(rid=rid,url=url,label=label))

    tb_map={}
    for pidx,para in enumerate(doc.paragraphs):
        for txbx in para._element.iter(qn("w:txbxContent")):
            t="".join((x.text or "") for x in txbx.iter(qn("w:t"))).strip()
            if t and pidx not in tb_map: tb_map[pidx]=t

    def _proc(paras,source="body"):
        nonlocal cur_sec,fid
        plist=list(paras)
        for pidx,para in enumerate(plist):
            full=_para_text(para).strip()

            if not full and source=="body" and pidx in tb_map:
                tb=tb_map[pidx]; sec=_detect_sec(tb,"Heading")
                if sec: cur_sec=sec
                fields.append(dict(id=f"f{fid}",text=tb,originalText=tb,
                    section=dict(cur_sec),isHeader=bool(sec),paraIndex=pidx,
                    source=source,runIndices=None,format="textbox",
                    isBold=True,hasHyperlink=False,isTextbox=True))
                fid+=1; continue

            if not full:
                wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                a_ns="http://schemas.openxmlformats.org/drawingml/2006/main"
                rects=[]
                for sp in para._element.iter(f"{{{wps}}}wsp"):
                    spx=etree.tostring(sp).decode()
                    geom=re.findall(r'prstGeom prst="([^"]+)"',spx)
                    ext=re.findall(r'<a:ext cx="(\d+)" cy="(\d+)"',spx)
                    if geom and "rect" in geom[0] and ext:
                        rects.append(dict(cx=int(ext[0][0]),cy=int(ext[0][1])))
                if len(rects)>=2 and rects[0]["cy"]==rects[1]["cy"]:
                    bg_w,fg_w=rects[0]["cx"],rects[1]["cx"]
                    if fg_w<=bg_w:
                        pct=round(fg_w/bg_w*100) if bg_w>0 else 0
                        lbl=""
                        try:
                            if pidx>0: lbl=_para_text(plist[pidx-1]).strip()
                        except: pass
                        fields.append(dict(id=f"f{fid}",text=f"{pct}%",
                            originalText=f"{pct}%",section=dict(cur_sec),
                            isHeader=False,type="progress-bar",paraIndex=pidx,
                            source=source,runIndices=None,
                            format=f"📊 {lbl or 'Skill'}",isBold=False,
                            hasHyperlink=False,barBgWidth=bg_w,
                            barFgWidth=fg_w,barLabel=lbl or "Skill"))
                        fid+=1
                continue

            sname=para.style.name if para.style else None
            sec=_detect_sec(full,sname)
            if sec:
                cur_sec=sec
                fields.append(dict(id=f"f{fid}",text=full,originalText=full,
                    section=dict(cur_sec),isHeader=True,paraIndex=pidx,
                    source=source,runIndices=None,format="section-header",
                    isBold=False,hasHyperlink=False))
                fid+=1; continue

            groups,all_r=_run_groups(para)
            has_hl=any(r["is_hl"] for r in all_r)

            # Collect inline hyperlink URLs from w:hyperlink elements in this para
            inline_links=[]
            seen_hl_rids=set()
            r_ns_hl="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            for hl_elem in para._element.findall(qn("w:hyperlink")):
                rid=hl_elem.get(f"{{{r_ns_hl}}}id") or ""
                # Some hyperlinks have no rId — mailto: or anchor-only links
                if rid:
                    if rid in seen_hl_rids: continue
                    seen_hl_rids.add(rid)
                    url=rel_map.get(rid,"")
                else:
                    # Try w:anchor (used for bookmarks) or inline URI on some converters
                    url=hl_elem.get(qn("w:anchor")) or ""
                if not url: continue
                hl_text="".join((t.text or "") for t in hl_elem.iter(qn("w:t"))).strip()
                # Also detect mailto from the display text if URL lookup failed
                if not url and re.search(r"[\w.+-]+@[\w.-]+\.\w+", hl_text):
                    url="mailto:"+re.search(r"[\w.+-]+@[\w.-]+\.\w+", hl_text).group()
                label=("LinkedIn" if "linkedin" in url.lower()
                       else "GitHub" if "github"   in url.lower()
                       else "Mail"   if url.lower().startswith("mailto:")
                       else "Link")
                inline_links.append(dict(rId=rid or "__noid__",url=url,text=hl_text,label=label))

            # If no mailto link found yet but field text contains an email address,
            # surface it as a synthetic "Mail" link so user can edit it
            import re as _re2
            if not any(lnk["label"]=="Mail" for lnk in inline_links):
                email_match = _re2.search(r'[\w.+-]+@[\w-]+\.\w+', full)
                if email_match:
                    # Check if this email has a hyperlink rId we may have missed (mailto: now included)
                    # If not found in inline_links, add a sentinel so the editor shows a Mail row
                    has_mail_hl = any(lnk["url"].lower().startswith("mailto:") for lnk in inline_links)
                    if not has_mail_hl:
                        inline_links.append(dict(rId="__plain_email__",
                            url="mailto:"+email_match.group(),
                            text=email_match.group(), label="Mail"))

            if len(groups)<=1:
                ft=_classify(full,cur_sec["key"],fid)
                fields.append(dict(id=f"f{fid}",text=full,originalText=full,
                    section=dict(cur_sec),isHeader=False,type=ft,
                    paraIndex=pidx,source=source,
                    runIndices=groups[0]["indices"] if groups else None,
                    format=groups[0]["fmt"] if groups else "regular",
                    isBold=groups[0]["is_bold"] if groups else False,
                    hasHyperlink=has_hl,
                    inlineLinks=inline_links if inline_links else None))
                fid+=1
            else:
                for g in groups:
                    t=g["text"].strip()
                    if not t: continue
                    ft=_classify(t,cur_sec["key"],fid)
                    fields.append(dict(id=f"f{fid}",text=t,originalText=t,
                        section=dict(cur_sec),isHeader=False,type=ft,
                        paraIndex=pidx,source=source,runIndices=g["indices"],
                        format=g["fmt"],isBold=g["is_bold"],
                        hasHyperlink=g.get("has_hl",False),
                        inlineLinks=inline_links if inline_links else None))
                    fid+=1

            if source=="body" and pidx in icon_links:
                for il in icon_links[pidx]:
                    fields.append(dict(id=f"f{fid}",text=il["url"],
                        originalText=il["url"],section=dict(cur_sec),
                        isHeader=False,type="icon-link",paraIndex=pidx,
                        source=source,runIndices=None,
                        format=f"🔗 {il['label']}",isBold=False,
                        hasHyperlink=True,linkRId=il["rid"],linkLabel=il["label"]))
                    fid+=1

    _proc(doc.paragraphs,"body")

    seen_fp=set()
    for ti,table in enumerate(doc.tables):
        fp="|||".join(
            "".join((te.text or "") for te in tc.iter(qn("w:t"))).strip()
            for tr in table._tbl.findall(qn("w:tr"))
            for tc in tr.findall(qn("w:tc")))
        if fp in seen_fp: continue
        seen_fp.add(fp)
        for ri2,tr in enumerate(table._tbl.findall(qn("w:tr"))):
            for ci,tc in enumerate(tr.findall(qn("w:tc"))):
                cell=next((c for c in table.rows[ri2].cells if c._tc is tc),None)
                if cell: _proc(cell.paragraphs,f"table-{ti}-{ri2}-{ci}")
                else:
                    from docx.text.paragraph import Paragraph as _P
                    _proc([_P(p,doc.element.body) for p in tc.findall(qn("w:p"))],
                          f"table-{ti}-{ri2}-{ci}")

    img_map={r.rId:r.target_ref for r in doc.part.rels.values()
             if "image" in str(r.reltype).lower()}
    seen_img=set()
    for pidx,para in enumerate(doc.paragraphs):
        a_ns="http://schemas.openxmlformats.org/drawingml/2006/main"
        for blip in para._element.iter(f"{{{a_ns}}}blip"):
            rid=blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rid or rid in seen_img or rid not in img_map: continue
            seen_img.add(rid); tgt=img_map[rid]
            fields.append(dict(id=f"f{fid}",
                text=f"📷 Image ({tgt.split('/')[-1]})",
                originalText=f"📷 Image ({tgt.split('/')[-1]})",
                section=dict(key="header",label="Header",color="#c026d3"),
                isHeader=False,type="image",paraIndex=pidx,source="body",
                runIndices=None,format="image",isBold=False,hasHyperlink=False,
                imageRId=rid,imageTarget=tgt))
            fid+=1

    fields = _classify_sections_llm(fields)
    return fields


def _source_paras(doc,source):
    if source=="body": return list(doc.paragraphs)
    if source.startswith("table-"):
        parts=source.split("-")
        try:
            ti,ri,ci=int(parts[1]),int(parts[2]),int(parts[3])
            tbl=doc.tables[ti]; trs=tbl._tbl.findall(qn("w:tr"))
            if ri<len(trs):
                tcs=trs[ri].findall(qn("w:tc"))
                if ci<len(tcs):
                    tc=tcs[ci]
                    cell=next((c for c in tbl.rows[ri].cells if c._tc is tc),None)
                    if cell: return list(cell.paragraphs)
                    from docx.text.paragraph import Paragraph as _P
                    return [_P(p,doc.element.body) for p in tc.findall(qn("w:p"))]
        except Exception as e: print(f"[source_paras] {e}")
    return list(doc.paragraphs)


# ════════════════════════════════════════════════════════════════════
#  SURGICAL TEXT REPLACEMENT
# ════════════════════════════════════════════════════════════════════

def _set_t(r_elem, text):
    ts=r_elem.findall(qn("w:t"))
    if ts:
        ts[0].text=text; ts[0].set(qn("xml:space"),"preserve")
        for t in ts[1:]: t.text=""
    elif text:
        rpr=r_elem.find(qn("w:rPr"))
        te=r_elem.makeelement(qn("w:t"),{qn("xml:space"):"preserve"})
        te.text=text
        if rpr is not None: rpr.addnext(te)
        else: r_elem.insert(0,te)

def _del_r(r_elem):
    p=r_elem.getparent()
    if p is not None: p.remove(r_elem)

def _strip_sp(r_elem):
    rpr=r_elem.find(qn("w:rPr"))
    if rpr is not None:
        sp=rpr.find(qn("w:spacing"))
        if sp is not None: rpr.remove(sp)

def _fix_tabs(para,all_r,valid):
    for idx in [min(valid)-1, max(valid)+1]:
        if 0<=idx<len(all_r):
            ri=all_r[idx]
            if ri["has_tab"] and not ri["text"].strip():
                tab=ri["elem"].find(qn("w:tab"))
                if tab is not None:
                    ri["elem"].remove(tab)
                    te=ri["elem"].makeelement(qn("w:t"),{qn("xml:space"):"preserve"})
                    te.text=" "; ri["elem"].append(te)
                    ppr=para._element.find(qn("w:pPr"))
                    if ppr is not None:
                        tabs=ppr.find(qn("w:tabs"))
                        if tabs is not None: ppr.remove(tabs)

def _repl_by_idx(para,run_indices,new_text):
    all_r=_get_runs(para)
    valid=[i for i in run_indices if i<len(all_r)]
    if not valid: return False
    _fix_tabs(para,all_r,valid)
    _set_t(all_r[valid[0]]["elem"],new_text); _strip_sp(all_r[valid[0]]["elem"])
    for idx in valid[1:]:
        ri=all_r[idx]
        if idx==valid[-1] and ri["text"] and not ri["text"].strip():
            _strip_sp(ri["elem"]); continue
        _del_r(ri["elem"])
    return True

def _repl_by_search(para,old_text,new_text):
    all_r=_get_runs(para); full=""; cmap=[]
    for i,ri in enumerate(all_r):
        for j in range(len(ri["text"])): cmap.append((i,j))
        full+=ri["text"]
    needle=old_text.strip(); pos=full.find(needle)
    if pos==-1: return False
    end=pos+len(needle)
    if end>len(cmap): return False
    sr,sc=cmap[pos]; er,ec=cmap[end-1]
    if sr==er:
        orig=all_r[sr]["text"]
        _set_t(all_r[sr]["elem"],orig[:sc]+new_text+orig[ec+1:]); _strip_sp(all_r[sr]["elem"])
    else:
        _set_t(all_r[sr]["elem"],all_r[sr]["text"][:sc]+new_text); _strip_sp(all_r[sr]["elem"])
        for k in range(sr+1,er): _del_r(all_r[k]["elem"])
        tail=all_r[er]["text"][ec+1:]
        if tail.strip(): _set_t(all_r[er]["elem"],tail); _strip_sp(all_r[er]["elem"])
        else: _del_r(all_r[er]["elem"])
    return True


# ════════════════════════════════════════════════════════════════════
#  SURGICAL FORMATTING
# ════════════════════════════════════════════════════════════════════

def _fmt_run(r_elem,run_obj,fmt):
    from docx.shared import Pt,RGBColor
    from docx.text.run import Run as _Run
    run=run_obj
    if run is None:
        try: run=_Run(r_elem,None)
        except: return

    if fmt.bold          is not None: run.bold      = fmt.bold
    if fmt.italic        is not None: run.italic    = fmt.italic
    if fmt.underline     is not None: run.underline = fmt.underline

    if fmt.strikethrough is not None:
        rpr=r_elem.find(qn("w:rPr"))
        if rpr is None:
            rpr=r_elem.makeelement(qn("w:rPr"),{}); r_elem.insert(0,rpr)
        strike=rpr.find(qn("w:strike"))
        if fmt.strikethrough:
            if strike is None: rpr.append(rpr.makeelement(qn("w:strike"),{}))
        else:
            if strike is not None: rpr.remove(strike)

    if fmt.font_color:
        h=fmt.font_color.lstrip("#")
        if len(h)==6:
            try: run.font.color.rgb=RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
            except Exception as e: print(f"[fmt] color: {e}")

    if fmt.font_name: run.font.name=fmt.font_name
    if fmt.font_size: run.font.size=Pt(float(fmt.font_size))

def _fmt_para(para,fmt):
    p_elem = para._element
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = p_elem.makeelement(qn("w:pPr"), {})
        p_elem.insert(0, pPr)

    if fmt.alignment:
        jc_map = {"left":"left","center":"center","right":"right","justify":"both"}
        jc_val = jc_map.get(fmt.alignment)
        if jc_val:
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = pPr.makeelement(qn("w:jc"), {})
                rpr_in_ppr = pPr.find(qn("w:rPr"))
                if rpr_in_ppr is not None:
                    pPr.insert(list(pPr).index(rpr_in_ppr), jc)
                else:
                    pPr.append(jc)
            jc.set(qn("w:val"), jc_val)

    if fmt.line_spacing is not None:
        try:
            multiplier = float(fmt.line_spacing)
            line_val = str(int(round(multiplier * 240)))
            sp = pPr.find(qn("w:spacing"))
            if sp is None:
                sp = pPr.makeelement(qn("w:spacing"), {})
                rpr_in_ppr = pPr.find(qn("w:rPr"))
                if rpr_in_ppr is not None:
                    pPr.insert(list(pPr).index(rpr_in_ppr), sp)
                else:
                    pPr.append(sp)
            sp.set(qn("w:line"), line_val)
            sp.set(qn("w:lineRule"), "auto")
            if sp.get(qn("w:lineRule")) == "exact":
                sp.set(qn("w:lineRule"), "auto")
        except Exception as e:
            print(f"[fmt] line_spacing error: {e}")

    if fmt.is_bullet is not None or fmt.is_numbered is not None:
        pPr=para._element.find(qn("w:pPr"))
        if pPr is None:
            pPr=para._element.makeelement(qn("w:pPr"),{}); para._element.insert(0,pPr)
        numPr=pPr.find(qn("w:numPr"))
        want=bool(fmt.is_bullet) or bool(fmt.is_numbered)
        if want:
            if numPr is None:
                numPr=pPr.makeelement(qn("w:numPr"),{}); pPr.insert(0,numPr)
            for tag,val in [(qn("w:ilvl"),"0"),(qn("w:numId"),"2" if fmt.is_numbered else "1")]:
                el=numPr.find(tag)
                if el is None: el=numPr.makeelement(tag,{}); numPr.append(el)
                el.set(qn("w:val"),val)
        else:
            if numPr is not None: pPr.remove(numPr)

def _apply_fmt(para, all_r, run_indices, fmt):
    bold_targets  = all_r
    other_targets = ([all_r[i] for i in run_indices if i < len(all_r)]
                     if run_indices else all_r)

    for ri in bold_targets:
        if not ri["text"].strip(): continue
        try:
            if fmt.bold is not None:
                _fmt_run(ri["elem"], ri.get("run"), FormatOptions(bold=fmt.bold))
        except Exception as e:
            print(f"[fmt] bold: {e}")

    for ri in other_targets:
        try:
            partial = FormatOptions(
                italic=fmt.italic, underline=fmt.underline,
                strikethrough=fmt.strikethrough, font_color=fmt.font_color,
                font_name=fmt.font_name, font_size=fmt.font_size,
            )
            _fmt_run(ri["elem"], ri.get("run"), partial)
        except Exception as e:
            print(f"[fmt] run: {e}")

    try: _fmt_para(para, fmt)
    except Exception as e: print(f"[fmt] para: {e}")


def apply_edit(doc_path,old_text,new_text,para_index,source,run_indices,output_path,formatting=None):
    doc=Document(doc_path); replaced=False

    def _try(para,ridx):
        nonlocal replaced
        ok=(_repl_by_idx(para,ridx,new_text) if ridx else False) or _repl_by_search(para,old_text,new_text)
        if ok:
            replaced=True
            if formatting:
                ar=_get_runs(para); _apply_fmt(para,ar,ridx,formatting)
        return ok

    paras=_source_paras(doc,source)
    if 0<=para_index<len(paras): _try(paras[para_index],run_indices)
    if not replaced:
        for p in paras:
            if old_text.strip() in _para_text(p):
                if _try(p,None): break
    if not replaced:
        all_p=list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: all_p.extend(cell.paragraphs)
        for p in all_p:
            if old_text.strip() in _para_text(p):
                if _try(p,None): break

    if replaced: doc.save(output_path)
    return replaced


def _edit_textbox(doc_path,old_text,new_text,output_path):
    doc=Document(doc_path); ok=False
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        t="".join((x.text or "") for x in txbx.iter(qn("w:t"))).strip()
        if old_text.strip()!=t: continue
        ts=[x for x in txbx.iter(qn("w:t")) if x.text and x.text.strip()]
        if ts:
            ts[0].text=new_text.strip()
            for x in ts[1:]: x.text=""
            ok=True
    if ok: doc.save(output_path)
    return ok

def _edit_link(doc_path,rid,new_url,output_path):
    with zipfile.ZipFile(doc_path,"r") as zi:
        with zipfile.ZipFile(output_path,"w") as zo:
            for item in zi.infolist():
                data=zi.read(item.filename)
                if item.filename=="word/_rels/document.xml.rels":
                    root=etree.fromstring(data); found=False
                    for rel in root:
                        if rel.get("Id")==rid: rel.set("Target",new_url); found=True; break
                    if not found: return False
                    data=etree.tostring(root,xml_declaration=True,encoding="UTF-8",standalone=True)
                zo.writestr(item,data)
    return True


# ════════════════════════════════════════════════════════════════════
#  ADD PARAGRAPH
# ════════════════════════════════════════════════════════════════════

def add_paragraph_after(doc_path, after_idx, source, segments, output_path,
                        inherit_list_format=True):
    doc=Document(doc_path); paras=_source_paras(doc,source)
    if not (0<=after_idx<len(paras)): return False
    ref=paras[after_idx]; ref_elem=ref._element

    new_p = ref_elem.makeelement(qn("w:p"), {})
    pPr = ref_elem.find(qn("w:pPr"))

    if pPr is not None:
        clean_pPr = deepcopy(pPr)
        rpr_in_ppr = clean_pPr.find(qn("w:rPr"))
        if rpr_in_ppr is not None: clean_pPr.remove(rpr_in_ppr)

        sp = clean_pPr.find(qn("w:spacing"))
        if sp is not None and sp.get(qn("w:lineRule")) == "exact":
            sp.attrib.pop(qn("w:lineRule"), None)
            sp.attrib.pop(qn("w:line"), None)

        num = clean_pPr.find(qn("w:numPr"))
        if num is not None and not inherit_list_format:
            clean_pPr.remove(num)

        new_p.append(clean_pPr)

    BULLET_CHARS={'•','◦','▪','▸','–','-','*'}
    ref_r=_get_runs(ref); bold_rpr=normal_rpr=None
    for ri in ref_r:
        txt=ri["text"].strip()
        if not txt: continue
        if len(txt)==1 and (txt in BULLET_CHARS or ord(txt)>127): continue
        rpr=ri["elem"].find(qn("w:rPr"))
        if rpr is None: continue
        fe=rpr.find(qn("w:rFonts"))
        sz=rpr.find(qn("w:sz"))
        if ri["bold"] and bold_rpr is None: bold_rpr=deepcopy(rpr)
        elif not ri["bold"] and normal_rpr is None: normal_rpr=deepcopy(rpr)
        if bold_rpr and normal_rpr: break

    for _t in (bold_rpr, normal_rpr):
        if _t is None: continue
        _sz = _t.find(qn("w:sz"))
        if _sz is not None and _sz.get(qn("w:val")):
            _szcs = _t.find(qn("w:szCs"))
            if _szcs is None:
                _szcs = _t.makeelement(qn("w:szCs"), {}); _t.append(_szcs)
            _szcs.set(qn("w:val"), _sz.get(qn("w:val")))

    if bold_rpr is None and normal_rpr is not None:
        bold_rpr=deepcopy(normal_rpr)
        if bold_rpr.find(qn("w:b")) is None: bold_rpr.insert(0,bold_rpr.makeelement(qn("w:b"),{}))
    elif normal_rpr is None and bold_rpr is not None:
        normal_rpr=deepcopy(bold_rpr)
        b=normal_rpr.find(qn("w:b"))
        if b is not None: normal_rpr.remove(b)

    def _clean_rpr(rpr):
        if rpr is None: return None
        rpr=deepcopy(rpr)
        for tag in (qn("w:spacing"),qn("w:kern"),qn("w:rStyle"),
                    qn("w:color"),qn("w:highlight"),qn("w:shd"),
                    qn("w:caps"),qn("w:smallCaps")):
            el=rpr.find(tag)
            if el is not None: rpr.remove(el)
        return rpr

    def _apply_seg_props(rpr, seg):
        import re as _re
        for tag in (qn("w:i"),qn("w:iCs")):
            el=rpr.find(tag)
            if getattr(seg,'italic',False):
                if el is None: rpr.append(rpr.makeelement(tag,{}))
            else:
                if el is not None: rpr.remove(el)
        u=rpr.find(qn("w:u"))
        if getattr(seg,'underline',False):
            if u is None: u=rpr.makeelement(qn("w:u"),{qn("w:val"):"single"}); rpr.append(u)
        else:
            if u is not None: rpr.remove(u)
        st=rpr.find(qn("w:strike"))
        if getattr(seg,'strike',False):
            if st is None: rpr.append(rpr.makeelement(qn("w:strike"),{}))
        else:
            if st is not None: rpr.remove(st)
        col=getattr(seg,'color','')
        if col:
            m=_re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)',col)
            hex_col=None
            if m: hex_col="{:02X}{:02X}{:02X}".format(int(m.group(1)),int(m.group(2)),int(m.group(3)))
            elif _re.match(r'^#?[0-9a-fA-F]{6}$',col): hex_col=col.lstrip("#").upper()
            if hex_col:
                c_el=rpr.find(qn("w:color"))
                if c_el is None: c_el=rpr.makeelement(qn("w:color"),{}); rpr.append(c_el)
                c_el.set(qn("w:val"),hex_col)
        else:
            c_el=rpr.find(qn("w:color"))
            if c_el is not None: rpr.remove(c_el)
        ff=getattr(seg,'fontFamily','')
        if ff:
            fe=rpr.find(qn("w:rFonts"))
            if fe is None: fe=rpr.makeelement(qn("w:rFonts"),{}); rpr.insert(0,fe)
            for a in (qn("w:ascii"),qn("w:hAnsi"),qn("w:cs")): fe.set(a,ff)
        fs=getattr(seg,'fontSize','')
        if fs:
            mp=_re.match(r'([\d.]+)pt',fs); mx=_re.match(r'([\d.]+)px',fs)
            pt=float(mp.group(1)) if mp else (float(mx.group(1))*0.75 if mx else None)
            if pt:
                half=str(int(pt*2))
                for tag in (qn("w:sz"),qn("w:szCs")):
                    sz=rpr.find(tag)
                    if sz is None: sz=rpr.makeelement(tag,{}); rpr.append(sz)
                    sz.set(qn("w:val"),half)

    def _mkr(text, seg):
        run=new_p.makeelement(qn("w:r"),{})
        tmpl=_clean_rpr(bold_rpr if seg.bold else normal_rpr)
        if tmpl is not None: run.append(tmpl)
        rpr=run.find(qn("w:rPr"))
        if rpr is None: rpr=run.makeelement(qn("w:rPr"),{}); run.insert(0,rpr)
        _apply_seg_props(rpr, seg)
        te=run.makeelement(qn("w:t"),{qn("xml:space"):"preserve"})
        te.text=text; run.append(te)
        return run

    MC_ALT       = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
    W_SYM        = qn("w:sym")
    W_DRAW       = qn("w:drawing")
    BULLET_CHARS = {'•','◦','▪','▸','–','–','•','-','*','·','○','□','▶','›','‣'}

    def _is_bullet_run(run_elem):
        txt = "".join(t.text or "" for t in run_elem.findall(qn("w:t")))
        stripped = txt.strip()
        has_special = (
            len(run_elem.findall('.//' + MC_ALT)) > 0 or
            len(run_elem.findall('.//' + W_SYM))  > 0 or
            len(run_elem.findall('.//' + W_DRAW)) > 0
        )
        is_bullet_char = (
            stripped and len(stripped) <= 3
            and all(c in BULLET_CHARS or ord(c) > 127 for c in stripped)
        )
        return has_special or is_bullet_char

    def _para_is_bullet_only(p_elem):
        has_bullet = False
        for ch in p_elem:
            if ch.tag != qn("w:r"): continue
            if _is_bullet_run(ch): has_bullet = True
            else:
                txt = "".join(t.text or "" for t in ch.findall(qn("w:t"))).strip()
                if txt: return False
        return has_bullet

    leading_bullet_runs = []
    found_text_run = False
    for ch in list(ref_elem):
        if ch.tag != qn("w:r"): continue
        stripped = "".join(t.text or "" for t in ch.findall(qn("w:t"))).strip()
        if not found_text_run and _is_bullet_run(ch):
            leading_bullet_runs.append(deepcopy(ch))
        else:
            if stripped: found_text_run = True

    W_TC  = qn("w:tc"); W_TR  = qn("w:tr"); W_TBL = qn("w:tbl")

    parent_tc  = ref_elem.getparent()
    parent_tr  = parent_tc.getparent() if parent_tc is not None else None
    in_table   = (parent_tc is not None and parent_tc.tag == W_TC and
                  parent_tr is not None and parent_tr.tag == W_TR)

    if in_table and not leading_bullet_runs:
        new_tr = deepcopy(parent_tr)
        cells = parent_tr.findall(W_TC)
        new_cells = new_tr.findall(W_TC)
        ref_cell_idx = None
        for i, tc in enumerate(cells):
            for p in tc.findall('.//' + qn("w:p")):
                if p is ref_elem: ref_cell_idx = i; break
            if ref_cell_idx is not None: break

        if ref_cell_idx is not None:
            new_text_cell = new_cells[ref_cell_idx]
            for old_p in new_text_cell.findall(qn("w:p")): new_text_cell.remove(old_p)
            for seg in segments:
                if seg.text: new_p.append(_mkr(seg.text, seg))
            new_text_cell.append(new_p)
            parent_tr.addnext(new_tr)
            doc.save(output_path)
            return True

    preceding_bullet_para = None
    prev_elem = ref_elem.getprevious()
    if not leading_bullet_runs and prev_elem is not None and prev_elem.tag == qn("w:p"):
        if _para_is_bullet_only(prev_elem):
            preceding_bullet_para = deepcopy(prev_elem)

    for seg in segments:
        if seg.text: new_p.append(_mkr(seg.text, seg))

    if preceding_bullet_para is not None:
        ref_elem.addnext(new_p)
        ref_elem.addnext(preceding_bullet_para)
    else:
        ppr_pos = None
        for i, ch in enumerate(list(new_p)):
            if ch.tag == qn("w:pPr"): ppr_pos = i; break
        insert_at = (ppr_pos + 1) if ppr_pos is not None else 0
        for br in reversed(leading_bullet_runs):
            new_p.insert(insert_at, br)
        ref_elem.addnext(new_p)

    doc.save(output_path); return True


# ════════════════════════════════════════════════════════════════════
#  PYDANTIC MODELS
# ════════════════════════════════════════════════════════════════════

class SegmentItem(BaseModel):
    text:       str
    bold:       bool = False
    italic:     bool = False
    underline:  bool = False
    strike:     bool = False
    color:      str  = ""
    fontFamily: str  = ""
    fontSize:   str  = ""

class FormatOptions(BaseModel):
    bold:          Optional[bool]  = None
    italic:        Optional[bool]  = None
    underline:     Optional[bool]  = None
    strikethrough: Optional[bool]  = None
    font_color:    Optional[str]   = None
    font_name:     Optional[str]   = None
    font_size:     Optional[float] = None
    alignment:     Optional[str]   = None
    line_spacing:  Optional[float] = None
    is_bullet:     Optional[bool]  = None
    is_numbered:   Optional[bool]  = None

class EditRequest(BaseModel):
    session_id:  str
    field_id:    str = ""
    old_text:    str
    new_text:    str
    para_index:  int = -1
    source:      str = "body"
    run_indices: Optional[List[int]]     = None
    link_rid:    Optional[str]           = None
    is_textbox:  bool                    = False
    set_bold:    Optional[bool]          = None
    formatting:  Optional[FormatOptions] = None

class TextSegment(BaseModel):
    text: str; bold: bool = False

class AddLineRequest(BaseModel):
    session_id:          str
    after_para_index:    int
    source:              str = "body"
    segments:            List[SegmentItem] = []
    after_field_text:    Optional[str] = None
    inherit_list_format: bool = True

class DeleteRequest(BaseModel):
    session_id: str; para_index: int; field_text: Optional[str] = None

class EditBarRequest(BaseModel):
    session_id: str; source: str; para_index: int; new_pct: int

# ── NEW: write a block of plain text to a blank page ─────────────────
class WritePageBlockRequest(BaseModel):
    session_id:        str
    after_para_index:  int   # index of the page-break paragraph; text goes right after it
    text:              str   # full text; newlines → separate paragraphs

# ── NEW: clear all content written to a blank page ────────────────────
class ClearPageBlockRequest(BaseModel):
    session_id:       str
    after_para_index: int   # page-break paragraph index; clears everything after it


# ════════════════════════════════════════════════════════════════════
#  ROUTES
# ════════════════════════════════════════════════════════════════════

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    fname=file.filename.lower()
    is_pdf=fname.endswith(".pdf"); is_docx=fname.endswith(".docx")
    if not is_pdf and not is_docx:
        raise HTTPException(400,"Only .docx or .pdf supported")
    if is_pdf and not (ILOVEPDF_OK or WORD_COM_OK or PDF2DOCX_OK or ADOBE_OK):
        raise HTTPException(400,"No PDF→DOCX converter available.")

    sid=str(uuid.uuid4())[:12]; d=UPLOAD_DIR/sid; d.mkdir(exist_ok=True)
    content=await file.read()

    if is_pdf:
        pdf_path=d/"original.pdf"; pdf_path.write_bytes(content)
        docx_path=d/"original.docx"
        ok,method=pdf_to_docx(pdf_path,docx_path)
        if not ok or not docx_path.exists():
            shutil.rmtree(d,ignore_errors=True)
            raise HTTPException(500,"PDF→DOCX failed.")
    else:
        (d/"original.docx").write_bytes(content)
        method="direct"

    shutil.copy(d/"original.docx",d/"current.docx")
    try: flatten_hyperlinks(str(d/"current.docx"))
    except Exception as e: print(f"[Upload] flatten: {e}")
    _patch_styles_gray(d/"current.docx")

    fields=extract_fields(str(d/"current.docx"))
    return dict(session_id=sid,filename=file.filename,
                fields=fields,
                field_count=len([f for f in fields if not f["isHeader"]]),
                converted_from_pdf=is_pdf,
                conversion_method=method)


@app.post("/edit")
async def edit_field(req: EditRequest):
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    tmp=d/"temp.docx"

    fmt=req.formatting
    if fmt is None and req.set_bold is not None:
        fmt=FormatOptions(bold=req.set_bold)
    elif fmt is not None and fmt.bold is None and req.set_bold is not None:
        fmt=fmt.copy(update=dict(bold=req.set_bold))

    if req.link_rid:
        ok=_edit_link(str(cur),req.link_rid,req.new_text.strip(),str(tmp))
        if ok: shutil.move(str(tmp),str(cur)); return dict(success=True,fields=extract_fields(str(cur)))
        if tmp.exists(): tmp.unlink()
        return dict(success=False,message="Link update failed")

    if req.is_textbox:
        ok=_edit_textbox(str(cur),req.old_text,req.new_text,str(tmp))
        if ok: shutil.move(str(tmp),str(cur)); return dict(success=True,fields=extract_fields(str(cur)))
        if tmp.exists(): tmp.unlink()
        return dict(success=False,message="Textbox update failed")

    ok=apply_edit(str(cur),req.old_text,req.new_text,
                  req.para_index,req.source,req.run_indices,str(tmp),formatting=fmt)
    if ok:
        shutil.move(str(tmp),str(cur))
        _patch_styles_gray(cur)
        return dict(success=True,fields=extract_fields(str(cur)))
    if tmp.exists(): tmp.unlink()
    return dict(success=False,message=f"Text not found: '{req.old_text[:50]}'")


class EditSegmentsRequest(BaseModel):
    session_id:  str
    field_id:    str = ""
    para_index:  int = -1
    source:      str = "body"
    segments:    List[SegmentItem]
    formatting:  Optional[FormatOptions] = None


@app.post("/edit-segments")
async def edit_segments(req: EditSegmentsRequest):
    print(f"[EditSegs] para={req.para_index} segs={len(req.segments)}")
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")

    try:
        doc=Document(str(cur))
        paras=_source_paras(doc, req.source)
        if not (0<=req.para_index<len(paras)):
            return dict(success=False, message="Paragraph index out of range")

        para=paras[req.para_index]
        p_elem=para._element

        pPr_elem = p_elem.find(qn("w:pPr"))
        if pPr_elem is not None:
            ppr_rpr = pPr_elem.find(qn("w:rPr"))
            if ppr_rpr is not None:
                for tag in (qn("w:rStyle"), qn("w:caps"), qn("w:smallCaps")):
                    el = ppr_rpr.find(tag)
                    if el is not None: ppr_rpr.remove(el)

        BULLET_CHARS={'•','◦','▪','▸','–','-','*'}
        all_r=_get_runs(para)
        normal_rpr=None; bold_rpr=None
        for ri in all_r:
            txt=ri["text"].strip()
            if not txt: continue
            if len(txt)==1 and (txt in BULLET_CHARS or ord(txt)>127): continue
            rpr=ri["elem"].find(qn("w:rPr"))
            if rpr is None: continue
            if ri["bold"] and bold_rpr is None:   bold_rpr=deepcopy(rpr)
            elif not ri["bold"] and normal_rpr is None: normal_rpr=deepcopy(rpr)
            if bold_rpr and normal_rpr: break

        if normal_rpr is None and bold_rpr is not None:
            normal_rpr=deepcopy(bold_rpr)
            b=normal_rpr.find(qn("w:b"))
            if b is not None: normal_rpr.remove(b)
            bcs=normal_rpr.find(qn("w:bCs"))
            if bcs is not None: normal_rpr.remove(bcs)
        if bold_rpr is None and normal_rpr is not None:
            bold_rpr=deepcopy(normal_rpr)
            if bold_rpr.find(qn("w:b")) is None:
                bold_rpr.insert(0, bold_rpr.makeelement(qn("w:b"),{}))

        def _clean_rpr(rpr):
            if rpr is None: return None
            rpr=deepcopy(rpr)
            for tag in (qn("w:spacing"),qn("w:kern"),qn("w:rStyle"),
                        qn("w:color"),qn("w:highlight"),qn("w:shd"),
                        qn("w:caps"),qn("w:smallCaps")):
                el=rpr.find(tag)
                if el is not None: rpr.remove(el)
            return rpr

        BULLET_CHARS_WIDE = {'•','◦','▪','▸','–','-','*','·','○','■','□','➢','✓','→'}
        preserved_leading = []
        all_children = list(p_elem)
        ALTERNATE_CONTENT_TAG = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
        W_SYM   = qn("w:sym")
        found_text_run = False
        for ch in all_children:
            if ch.tag == qn("w:r"):
                txt = "".join(t.text or "" for t in ch.findall(qn("w:t")))
                stripped = txt.strip()
                has_alt = ch.find(ALTERNATE_CONTENT_TAG) is not None
                has_sym = ch.find(W_SYM) is not None
                is_bullet_char = (stripped and len(stripped) <= 2
                    and all(c in BULLET_CHARS_WIDE or ord(c) > 127 for c in stripped))
                if not found_text_run and (has_alt or has_sym or is_bullet_char):
                    preserved_leading.append(deepcopy(ch))
                else:
                    if txt.strip():
                        found_text_run = True
            elif ch.tag == qn("w:hyperlink"):
                found_text_run = True

        structural_tab_runs = []
        for ch in all_children:
            if ch.tag != qn("w:r"): continue
            ch_txt  = "".join(t.text or "" for t in ch.findall(qn("w:t")))
            has_tab = ch.find(qn("w:tab")) is not None
            has_br  = ch.find(qn("w:br"))  is not None
            if (has_tab or has_br) and not ch_txt.strip():
                structural_tab_runs.append(deepcopy(ch))

        # ── Save hyperlink structure BEFORE clearing runs ────────────────
        # Map: rId → (hyperlink_text, original_elem_copy)
        saved_hyperlinks = []
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        for hl in p_elem.findall(qn("w:hyperlink")):
            rid = hl.get(f"{{{r_ns}}}id") or ""
            hl_text = "".join((t.text or "") for t in hl.iter(qn("w:t")))
            if rid and hl_text.strip():
                saved_hyperlinks.append({"rId": rid, "text": hl_text, "elem": deepcopy(hl)})

        for ch in list(p_elem):
            if ch.tag in (qn("w:r"), qn("w:hyperlink"),
                          qn("w:ins"), qn("w:del")):
                p_elem.remove(ch)

        pPr_elem2 = p_elem.find(qn("w:pPr"))
        insert_after = pPr_elem2 if pPr_elem2 is not None else None
        for br in preserved_leading:
            if insert_after is not None:
                idx_ia = list(p_elem).index(insert_after)
                p_elem.insert(idx_ia + 1, br)
                insert_after = br
            else:
                p_elem.insert(0, br)

        for seg in req.segments:
            if not seg.text: continue

            if seg.text == "	":
                if structural_tab_runs:
                    tab_elem = structural_tab_runs.pop(0)
                    p_elem.append(tab_elem)
                else:
                    tab_run = p_elem.makeelement(qn("w:r"), {})
                    tab_t   = tab_run.makeelement(qn("w:tab"), {})
                    tab_run.append(tab_t)
                    p_elem.append(tab_run)
                continue

            run=p_elem.makeelement(qn("w:r"),{})
            tmpl=_clean_rpr(bold_rpr if seg.bold else normal_rpr)
            if tmpl is not None: run.append(tmpl)

            rpr=run.find(qn("w:rPr"))
            if rpr is None:
                rpr=run.makeelement(qn("w:rPr"),{}); run.insert(0,rpr)

            st_elem=rpr.find(qn("w:strike"))
            if getattr(seg,'strike',False):
                if st_elem is None: rpr.append(rpr.makeelement(qn("w:strike"),{}))
            else:
                if st_elem is not None: rpr.remove(st_elem)

            for tag in (qn("w:i"),qn("w:iCs")):
                el=rpr.find(tag)
                if seg.italic:
                    if el is None: rpr.append(rpr.makeelement(tag,{}))
                else:
                    if el is not None: rpr.remove(el)

            u_elem=rpr.find(qn("w:u"))
            if seg.underline:
                if u_elem is None:
                    u_elem=rpr.makeelement(qn("w:u"),{qn("w:val"):"single"}); rpr.append(u_elem)
            else:
                if u_elem is not None: rpr.remove(u_elem)

            import re as _re
            if seg.color:
                hex_col=None
                m=_re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)',seg.color)
                if m: hex_col="{:02X}{:02X}{:02X}".format(int(m.group(1)),int(m.group(2)),int(m.group(3)))
                elif _re.match(r'^#?[0-9a-fA-F]{6}$',seg.color): hex_col=seg.color.lstrip("#").upper()
                if hex_col:
                    col_elem=rpr.find(qn("w:color"))
                    if col_elem is None:
                        col_elem=rpr.makeelement(qn("w:color"),{}); rpr.append(col_elem)
                    col_elem.set(qn("w:val"),hex_col)
            else:
                col_elem=rpr.find(qn("w:color"))
                if col_elem is not None: rpr.remove(col_elem)

            if seg.fontFamily:
                fonts_elem=rpr.find(qn("w:rFonts"))
                if fonts_elem is None:
                    fonts_elem=rpr.makeelement(qn("w:rFonts"),{}); rpr.insert(0,fonts_elem)
                for attr in (qn("w:ascii"),qn("w:hAnsi"),qn("w:cs")):
                    fonts_elem.set(attr,seg.fontFamily)

            if seg.fontSize:
                pt=None
                mp=_re.match(r'([\d.]+)pt',seg.fontSize)
                mx=_re.match(r'([\d.]+)px',seg.fontSize)
                if mp: pt=float(mp.group(1))
                elif mx: pt=float(mx.group(1))*0.75
                if pt:
                    half=str(int(pt*2))
                    for tag in (qn("w:sz"),qn("w:szCs")):
                        sz=rpr.find(tag)
                        if sz is None: sz=rpr.makeelement(tag,{}); rpr.append(sz)
                        sz.set(qn("w:val"),half)

            te=run.makeelement(qn("w:t"),{qn("xml:space"):"preserve"})
            te.text=seg.text; run.append(te)
            p_elem.append(run)

        for tab_elem in structural_tab_runs:
            p_elem.append(tab_elem)

        if req.formatting:
            try: _fmt_para(para, req.formatting)
            except Exception as e: print(f"[EditSegs] para fmt: {e}")

        # ── Restore hyperlink wrappers for text that still exists ────────
        # Build the full new paragraph text for substring matching
        if saved_hyperlinks:
            new_full = "".join(
                (t.text or "") for r in p_elem.findall(qn("w:r"))
                for t in r.findall(qn("w:t"))
            )
            for hl_info in saved_hyperlinks:
                hl_text = hl_info["text"]
                rid     = hl_info["rId"]
                if not hl_text.strip() or hl_text not in new_full:
                    continue  # text changed → link lost, that's intentional

                # Find the run(s) whose combined text matches hl_text
                runs = p_elem.findall(qn("w:r"))
                start_r = end_r = None
                accumulated = ""
                for i, r in enumerate(runs):
                    r_text = "".join((t.text or "") for t in r.findall(qn("w:t")))
                    if hl_text.startswith(accumulated + r_text):
                        if start_r is None: start_r = i
                        accumulated += r_text
                        if accumulated == hl_text:
                            end_r = i; break
                    else:
                        accumulated = ""; start_r = None

                if start_r is None or end_r is None:
                    continue

                # Build a new w:hyperlink element with the original rId
                new_hl = p_elem.makeelement(qn("w:hyperlink"), {})
                new_hl.set(f"{{{r_ns}}}id", rid)
                # Move matching runs into the hyperlink
                target_runs = runs[start_r:end_r + 1]
                first_run = target_runs[0]
                insert_pos = list(p_elem).index(first_run)
                for r in target_runs:
                    p_elem.remove(r)
                    new_hl.append(r)
                p_elem.insert(insert_pos, new_hl)
                print(f"[EditSegs] restored hyperlink rId={rid} text='{hl_text[:30]}'")

        tmp=d/"temp.docx"; doc.save(str(tmp))
        shutil.move(str(tmp),str(cur))
        _patch_styles_gray(cur)
        return dict(success=True, fields=extract_fields(str(cur)))

    except Exception as e:
        import traceback
        print(f"[EditSegs] error: {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))


@app.post("/add-line")
async def add_line(req: AddLineRequest):
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    if not any(s.text.strip() for s in req.segments):
        return dict(success=False,message="Text cannot be empty")
    tmp=d/"temp.docx"
    ok=add_paragraph_after(
        str(cur), req.after_para_index, req.source, req.segments, str(tmp),
        inherit_list_format=req.inherit_list_format,
    )
    if ok:
        shutil.move(str(tmp),str(cur))
        _patch_styles_gray(cur)
        return dict(success=True,fields=extract_fields(str(cur)))
    if tmp.exists(): tmp.unlink()
    return dict(success=False,message="Could not add line")


@app.post("/delete-line")
async def delete_line(req: DeleteRequest):
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    try:
        doc=Document(str(cur)); paras=doc.paragraphs
        if not (0<=req.para_index<len(paras)):
            return dict(success=False,message="Index out of range")
        paras[req.para_index]._element.getparent().remove(paras[req.para_index]._element)
        tmp=d/"temp.docx"; doc.save(str(tmp)); shutil.move(str(tmp),str(cur))
        return dict(success=True,fields=extract_fields(str(cur)))
    except Exception as e: return dict(success=False,message=str(e))


@app.post("/delete-field")
async def delete_field_ep(req: DeleteRequest):
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    try:
        doc=Document(str(cur)); paras=doc.paragraphs; target=None
        if 0<=req.para_index<len(paras): target=paras[req.para_index]
        elif req.field_text:
            for p in paras:
                if req.field_text.strip() in _para_text(p): target=p; break
        if target is None: return dict(success=False,message="Field not found")
        target._element.getparent().remove(target._element)
        tmp=d/"temp.docx"; doc.save(str(tmp)); shutil.move(str(tmp),str(cur))
        return dict(success=True,fields=extract_fields(str(cur)))
    except Exception as e: return dict(success=False,message=str(e))


@app.post("/replace-image")
async def replace_image(session_id: str, image_target: str, file: UploadFile=File(...)):
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    img=await file.read()
    if not img: return dict(success=False,message="Empty file")
    tgt=f"word/{image_target}"; tmp=d/"temp.docx"
    try:
        with zipfile.ZipFile(str(cur),"r") as zi:
            with zipfile.ZipFile(str(tmp),"w") as zo:
                for item in zi.infolist():
                    zo.writestr(item,img if item.filename==tgt else zi.read(item.filename))
        shutil.move(str(tmp),str(cur))
        return dict(success=True,fields=extract_fields(str(cur)))
    except Exception as e:
        if tmp.exists(): tmp.unlink()
        return dict(success=False,message=str(e))


@app.post("/edit-bar")
async def edit_bar(req: EditBarRequest):
    d=UPLOAD_DIR/req.session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    doc=Document(str(cur)); pct=max(1,min(100,req.new_pct))
    paras=_source_paras(doc,req.source)
    if req.para_index>=len(paras): return dict(success=False,message="Paragraph not found")
    para=paras[req.para_index]
    wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    a_ns="http://schemas.openxmlformats.org/drawingml/2006/main"
    rects=[]
    for sp in para._element.iter(f"{{{wps}}}wsp"):
        spx=etree.tostring(sp).decode()
        geom=re.findall(r'prstGeom prst="([^"]+)"',spx)
        exts=list(sp.iter(f"{{{a_ns}}}ext"))
        if geom and "rect" in geom[0] and exts: rects.append(exts)
    if len(rects)<2: return dict(success=False,message="No bar found")
    bg_w=int(rects[0][0].get("cx","0")); new_fg=max(1,int(bg_w*pct/100))
    for ext in rects[1]: ext.set("cx",str(new_fg))
    tmp=d/"temp.docx"; doc.save(str(tmp)); shutil.move(str(tmp),str(cur))
    return dict(success=True,fields=extract_fields(str(cur)))


@app.post("/reset/{session_id}")
async def reset(session_id: str):
    d=UPLOAD_DIR/session_id; orig=d/"original.docx"
    if not orig.exists(): raise HTTPException(404,"Session not found")
    shutil.copy(orig,d/"current.docx")
    try: flatten_hyperlinks(str(d/"current.docx"))
    except: pass
    for f in [d/"preview.pdf",d/"download.pdf"]:
        if f.exists(): f.unlink()
    return dict(success=True,fields=extract_fields(str(d/"current.docx")))


@app.get("/download/{session_id}")
async def download_docx(session_id: str):
    p=UPLOAD_DIR/session_id/"current.docx"
    if not p.exists(): raise HTTPException(404,"Not found")
    return FileResponse(str(p),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="resume_edited.docx")


@app.get("/download-pdf/{session_id}")
async def download_pdf(session_id: str):
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    pdf=d/"download.pdf"; ok=docx_to_pdf(cur,pdf)
    if not ok or not pdf.exists():
        raise HTTPException(500,"PDF conversion failed — install docx2pdf or LibreOffice")
    return FileResponse(str(pdf),media_type="application/pdf",filename="resume_edited.pdf")


@app.get("/preview-pdf/{session_id}")
async def preview_pdf(session_id: str):
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Session not found")
    pdf=d/"preview.pdf"; ok=docx_to_pdf(cur,pdf)
    if not ok or not pdf.exists():
        raise HTTPException(500,"Preview failed — install docx2pdf or LibreOffice")
    return Response(content=pdf.read_bytes(),media_type="application/pdf",
                    headers={"Content-Disposition":"inline",
                             "Cache-Control":"no-cache, no-store, must-revalidate"})


# ════════════════════════════════════════════════════════════════════
#  NEW: Write a block of plain text to a blank page
#       Inserts one paragraph per newline, right after the page-break
#       paragraph whose index is supplied by the frontend.
# ════════════════════════════════════════════════════════════════════

@app.post("/write-page-block")
async def write_page_block(req: WritePageBlockRequest):
    """
    Write multi-line text onto a blank page.
    Each newline in `req.text` becomes a separate paragraph.
    Blank lines are preserved as empty spacing paragraphs.
    Paragraphs are inserted immediately after `req.after_para_index`
    (the page-break paragraph added by /add-page).
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")

    try:
        doc = Document(str(cur))
        paras = doc.paragraphs

        if not (0 <= req.after_para_index < len(paras)):
            return dict(success=False, message="Paragraph index out of range")

        lines = req.text.split("\n")
        ref_elem = paras[req.after_para_index]._element
        insert_after = ref_elem

        # Auto-insert a page break every PAGE_BREAK_EVERY non-empty lines
        # so large content creates multiple progress-bar sections instead of
        # overflowing silently into an un-trackable visual page.
        PAGE_BREAK_EVERY = 50
        non_empty_count = 0

        for line_text in lines:
            # Insert auto page break BEFORE this line if threshold reached
            if non_empty_count > 0 and non_empty_count % PAGE_BREAK_EVERY == 0 and line_text.strip():
                pb_p = doc.element.body.makeelement(qn("w:p"), {})
                pb_r = doc.element.body.makeelement(qn("w:r"), {})
                pb_b = doc.element.body.makeelement(qn("w:br"), {})
                pb_b.set(qn("w:type"), "page")
                pb_r.append(pb_b)
                pb_p.append(pb_r)
                insert_after.addnext(pb_p)
                insert_after = pb_p

            new_p = doc.element.body.makeelement(qn("w:p"), {})
            if line_text.strip():
                run = new_p.makeelement(qn("w:r"), {})
                t   = run.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                t.text = line_text
                run.append(t)
                new_p.append(run)
                non_empty_count += 1

            insert_after.addnext(new_p)
            insert_after = new_p

        auto_breaks = non_empty_count // PAGE_BREAK_EVERY
        tmp = d / "temp.docx"
        doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        print(f"[WritePageBlock] wrote {len(lines)} lines, {auto_breaks} auto page-breaks inserted")
        return dict(success=True, fields=extract_fields(str(cur)))

    except Exception as e:
        import traceback
        print(f"[WritePageBlock] error: {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))

# ════════════════════════════════════════════════════════════════════
#  NEW: Clear all content written to a blank page
#       Removes every paragraph between the page-break paragraph and
#       the next page break (or end of document), then re-inserts a
#       small number of blank spacer paragraphs so the page still exists.
# ════════════════════════════════════════════════════════════════════

@app.post("/clear-page-block")
async def clear_page_block(req: ClearPageBlockRequest):
    """
    Delete all paragraphs written after the page-break paragraph.
    Stops at the next page-break or end of document.
    Restores 10 blank spacer paragraphs so the page still renders.

    Key fix: ref_elem is captured BEFORE removal so para-index shift
    after deletion can never break the spacer re-insertion.
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")

    try:
        doc = Document(str(cur))
        paras = doc.paragraphs

        if not (0 <= req.after_para_index < len(paras)):
            return dict(success=False, message=f"Paragraph index {req.after_para_index} out of range (doc has {len(paras)} paras)")

        # ── Capture the page-break element BEFORE any removal ──────────
        ref_elem = paras[req.after_para_index]._element

        # ── Collect paragraphs to remove ───────────────────────────────
        # Walk forward from after_para_index+1; stop at the next page break
        to_remove = []
        for i in range(req.after_para_index + 1, len(paras)):
            p = paras[i]
            has_page_break = any(
                br.get(qn("w:type")) == "page"
                for br in p._element.iter(qn("w:br"))
            )
            if has_page_break:
                break
            to_remove.append(p._element)

        if not to_remove:
            # Nothing to remove — page is already empty, just confirm success
            print(f"[ClearPageBlock] nothing to remove after para {req.after_para_index}")
            return dict(success=True, fields=extract_fields(str(cur)), removed=0)

        # ── Remove ─────────────────────────────────────────────────────
        for p_elem in to_remove:
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)

        # ── Re-insert blank spacers using the saved ref_elem ───────────
        body = doc.element.body
        insert_after = ref_elem
        for _ in range(10):
            new_p = body.makeelement(qn("w:p"), {})
            insert_after.addnext(new_p)
            insert_after = new_p

        tmp = d / "temp.docx"
        doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        removed_count = len(to_remove)
        print(f"[ClearPageBlock] removed {removed_count} paragraphs after para {req.after_para_index}")
        return dict(success=True, fields=extract_fields(str(cur)), removed=removed_count)

    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print(f"[ClearPageBlock] error: {e}\n{tb}")
        return dict(success=False, message=str(e))



# ════════════════════════════════════════════════════════════════════
#  NEW: Delete the entire blank page
#       Removes the page-break paragraph itself AND every paragraph
#       after it up to (but not including) the next page-break.
#       Uses the same ClearPageBlockRequest model (same fields).
# ════════════════════════════════════════════════════════════════════

@app.post("/delete-page-block")
async def delete_page_block(req: ClearPageBlockRequest):
    """
    Permanently remove a blank page from the document:
    - Deletes the page-break paragraph at after_para_index
    - Deletes all following paragraphs up to the next page-break (or end of doc)
    On success the frontend should reset pageWriterParaIdx to -1.
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")

    try:
        doc = Document(str(cur))
        paras = doc.paragraphs

        if not (0 <= req.after_para_index < len(paras)):
            return dict(success=False,
                        message=f"Paragraph index {req.after_para_index} out of range "
                                f"(doc has {len(paras)} paras)")

        # Include the page-break paragraph itself in the removal set
        to_remove = [paras[req.after_para_index]._element]

        # Collect all paragraphs that belong to this page
        for i in range(req.after_para_index + 1, len(paras)):
            p = paras[i]
            has_next_page_break = any(
                br.get(qn("w:type")) == "page"
                for br in p._element.iter(qn("w:br"))
            )
            if has_next_page_break:
                break          # stop before the next page break — don't remove it
            to_remove.append(p._element)

        for p_elem in to_remove:
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)

        tmp = d / "temp.docx"
        doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        print(f"[DeletePageBlock] removed {len(to_remove)} paras "
              f"(page break + content) at index {req.after_para_index}")
        return dict(success=True, fields=extract_fields(str(cur)), removed=len(to_remove))

    except Exception as e:
        import traceback
        print(f"[DeletePageBlock] error: {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))



@app.get("/get-page-breaks/{session_id}")
async def get_page_breaks(session_id: str):
    """Return the para index of every page-break paragraph in the document,
    sorted ascending. Used by the frontend to track and target individual pages."""
    d = UPLOAD_DIR / session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")
    try:
        doc = Document(str(cur))
        breaks = []
        for i, para in enumerate(doc.paragraphs):
            for br in para._element.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    breaks.append(i)
                    break  # one page-break per para is enough
        return dict(page_breaks=sorted(breaks))
    except Exception as e:
        return dict(page_breaks=[], error=str(e))



@app.get("/debug-hyperlinks/{session_id}")
async def debug_hyperlinks(session_id: str, para: int = 0):
    """Show all w:hyperlink elements and their resolved URLs for a paragraph."""
    d = UPLOAD_DIR / session_id; cur = d / "current.docx"
    if not cur.exists(): raise HTTPException(404, "Not found")
    doc = Document(str(cur))
    rel_map = {r.rId: r.target_ref for r in doc.part.rels.values()
               if "hyperlink" in str(r.reltype).lower()}
    result = []
    paras = doc.paragraphs
    if para < len(paras):
        p = paras[para]
        for hl in p._element.findall(qn("w:hyperlink")):
            ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            rid = hl.get(f"{{{ns}}}id") or hl.get(qn("w:id")) or ""
            url = rel_map.get(rid, "NOT IN REL_MAP")
            text = "".join((t.text or "") for t in hl.iter(qn("w:t")))
            result.append({"rId": rid, "url": url, "text": text})
    return {"para": para, "para_text": _para_text(paras[para]) if para < len(paras) else "",
            "rel_map_keys": list(rel_map.keys())[:10],
            "hyperlinks": result}



class RelinkRequest(BaseModel):
    session_id: str
    para_index: int
    source:     str = "body"
    text:       str   # display text to find and wrap
    url:        str   # target URL to link to


class EditHyperlinkTextRequest(BaseModel):
    session_id: str
    para_index: int
    source:     str = "body"
    r_id:       str   # relationship ID of the w:hyperlink to edit
    new_text:   str   # new display text for the hyperlink

@app.post("/edit-hyperlink-text")
async def edit_hyperlink_text(req: EditHyperlinkTextRequest):
    """
    Change the display text of a specific w:hyperlink element (found by rId).
    Only touches w:t nodes INSIDE that hyperlink — wrapper and relationship intact.
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")
    try:
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        doc  = Document(str(cur))
        paras = _source_paras(doc, req.source)
        if not (0 <= req.para_index < len(paras)):
            return dict(success=False, message="Paragraph index out of range")

        p_elem = paras[req.para_index]._element
        target_hl = None
        for hl in p_elem.iter(qn("w:hyperlink")):
            if hl.get(f"{{{r_ns}}}id") == req.r_id:
                target_hl = hl
                break

        if target_hl is None:
            return dict(success=False, message=f"Hyperlink rId={req.r_id} not found in paragraph")

        # Replace all text runs inside the hyperlink with the new text
        runs = target_hl.findall(qn("w:r"))
        if not runs:
            return dict(success=False, message="Hyperlink has no text runs")

        # Set first run to new text, clear the rest
        first_r = runs[0]
        t_elems = first_r.findall(qn("w:t"))
        if t_elems:
            t_elems[0].text = req.new_text
            t_elems[0].set(qn("xml:space"), "preserve")
            for t in t_elems[1:]: t.text = ""
        else:
            te = first_r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
            te.text = req.new_text
            first_r.append(te)

        for r in runs[1:]:
            for t in r.findall(qn("w:t")): t.text = ""

        tmp = d / "temp.docx"
        doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        print(f"[EditHyperlinkText] rId={req.r_id} → '{req.new_text[:40]}'")
        return dict(success=True, fields=extract_fields(str(cur)))

    except Exception as e:
        import traceback
        print(f"[EditHyperlinkText] error: {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))

@app.post("/relink-text")
async def relink_text(req: RelinkRequest):
    """
    Find `req.text` in the paragraph and wrap it in a new w:hyperlink element.
    Adds the URL to the document relationships and creates the XML wrapper.
    Works even when the hyperlink was previously destroyed by an edit.
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")
    try:
        import zipfile as _zf, io as _io, uuid as _uuid, re as _re
        from lxml import etree as _et

        doc = Document(str(cur))
        paras = _source_paras(doc, req.source)
        if not (0 <= req.para_index < len(paras)):
            return dict(success=False, message="Paragraph index out of range")

        para = paras[req.para_index]
        p_elem = para._element
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        # Build full text + char→run mapping from ALL runs (including inside hyperlinks)
        # Using iter() instead of findall() to traverse into w:hyperlink children
        full_text = ""
        run_spans = []  # (start, end, run_elem, parent_elem)
        for child in p_elem:
            if child.tag == qn("w:r"):
                r_text = "".join(t.text or "" for t in child.findall(qn("w:t")))
                run_spans.append((len(full_text), len(full_text) + len(r_text), child, p_elem))
                full_text += r_text
            elif child.tag == qn("w:hyperlink"):
                for r in child.findall(qn("w:r")):
                    r_text = "".join(t.text or "" for t in r.findall(qn("w:t")))
                    run_spans.append((len(full_text), len(full_text) + len(r_text), r, child))
                    full_text += r_text

        # Support partial/substring match — find the text anywhere in the paragraph
        pos = full_text.find(req.text)
        if pos == -1:
            # Try case-insensitive
            pos = full_text.lower().find(req.text.lower())
        if pos == -1:
            return dict(success=False,
                message=f"Text not found in paragraph: '{req.text[:40]}' | Para text: '{full_text[:80]}'")

        # Find overlapping runs (4-tuple: start, end, run_elem, parent_elem)
        end = pos + len(req.text)
        overlap = [(rs, re, r, par) for rs, re, r, par in run_spans if re > pos and rs < end]
        if not overlap:
            return dict(success=False, message="No runs found for the text range")

        # If the text is already inside a hyperlink, just update that hyperlink's URL
        unique_parents = set(id(par) for _, _, _, par in overlap)
        if len(unique_parents) == 1:
            parent = overlap[0][3]
            if parent.tag == qn("w:hyperlink"):
                # Already hyperlinked — update the relationship target
                old_rid = parent.get(f"{{{r_ns}}}id") or ""
                if old_rid:
                    # Save doc first then patch rels
                    doc.save(str(tmp))
                    buf2 = _io.BytesIO()
                    with _zf.ZipFile(str(tmp), "r") as zi:
                        with _zf.ZipFile(buf2, "w", _zf.ZIP_DEFLATED) as zo:
                            for item in zi.infolist():
                                data = zi.read(item.filename)
                                if item.filename == "word/_rels/document.xml.rels":
                                    root = _et.fromstring(data)
                                    for rel in root:
                                        if rel.get("Id") == old_rid:
                                            rel.set("Target", req.url)
                                            break
                                    data = _et.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                                zo.writestr(item, data)
                    buf2.seek(0)
                    with open(str(cur), "wb") as f_out: f_out.write(buf2.read())
                    _patch_styles_gray(cur)
                    print(f"[RelinkText] updated existing hyperlink {old_rid} → '{req.url}'")
                    return dict(success=True, fields=extract_fields(str(cur)))

        # Split first run if it starts before pos
        fs, fe, fr, fpar = overlap[0]
        if fs < pos:
            before_text = full_text[fs:pos]
            after_text  = full_text[pos:fe]
            for t in list(fr.findall(qn("w:t"))): fr.remove(t)
            if before_text:
                te = fr.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                te.text = before_text; fr.append(te)
            new_r = deepcopy(fr)
            for t in list(new_r.findall(qn("w:t"))): new_r.remove(t)
            if after_text:
                te = new_r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                te.text = after_text; new_r.append(te)
            idx = list(fpar).index(fr)
            fpar.insert(idx + 1, new_r)
            overlap[0] = (pos, fe, new_r, fpar)

        # Split last run if it ends after end
        ls, le, lr, lpar = overlap[-1]
        if le > end:
            split_text_in  = full_text[ls:end]
            split_text_out = full_text[end:le]
            for t in list(lr.findall(qn("w:t"))): lr.remove(t)
            if split_text_in:
                te = lr.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                te.text = split_text_in; lr.append(te)
            after_r = deepcopy(lr)
            for t in list(after_r.findall(qn("w:t"))): after_r.remove(t)
            if split_text_out:
                te = after_r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                te.text = split_text_out; after_r.append(te)
            idx = list(lpar).index(lr)
            lpar.insert(idx + 1, after_r)

        # Add new relationship to the document
        new_rid = f"rIdRL{_uuid.uuid4().hex[:8]}"
        tmp = d / "temp.docx"
        doc.save(str(tmp))

        # Patch the relationships XML to add the new rId → url
        buf = _io.BytesIO()
        with _zf.ZipFile(str(tmp), "r") as zi:
            with _zf.ZipFile(buf, "w", _zf.ZIP_DEFLATED) as zo:
                for item in zi.infolist():
                    data = zi.read(item.filename)
                    if item.filename == "word/_rels/document.xml.rels":
                        root = _et.fromstring(data)
                        new_rel = root.makeelement("Relationship", {
                            "Id":     new_rid,
                            "Type":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                            "Target": req.url,
                            "TargetMode": "External",
                        })
                        root.append(new_rel)
                        data = _et.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    zo.writestr(item, data)
        buf.seek(0)
        with open(str(tmp), "wb") as f_out: f_out.write(buf.read())

        # Now re-open and wrap the found runs in a w:hyperlink
        doc2 = Document(str(tmp))
        paras2 = _source_paras(doc2, req.source)
        p2 = paras2[req.para_index]._element
        runs2 = p2.findall(qn("w:r"))

        # Find the runs corresponding to our text (re-search in updated doc)
        full2 = ""; rspans2 = []
        for r in runs2:
            rt = "".join(t.text or "" for t in r.findall(qn("w:t")))
            rspans2.append((len(full2), len(full2) + len(rt), r))
            full2 += rt

        pos2 = full2.find(req.text)
        if pos2 == -1:
            return dict(success=False, message="Text not found after split")
        end2 = pos2 + len(req.text)
        wrap_runs = [r for rs, re, r in rspans2 if re > pos2 and rs < end2]

        if not wrap_runs:
            return dict(success=False, message="Could not identify runs to wrap")

        new_hl = p2.makeelement(qn("w:hyperlink"), {})
        new_hl.set(f"{{{r_ns}}}id", new_rid)
        insert_idx = list(p2).index(wrap_runs[0])
        for r in wrap_runs:
            p2.remove(r)
            new_hl.append(r)
        p2.insert(insert_idx, new_hl)

        doc2.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        print(f"[RelinkText] re-linked '{req.text[:30]}' → '{req.url}' as {new_rid}")
        return dict(success=True, fields=extract_fields(str(cur)))

    except Exception as e:
        import traceback
        print(f"[RelinkText] error: {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))


@app.get("/debug-sections/{session_id}")
async def debug_sections(session_id: str):
    """Show every field with its detected section and paraIndex — for diagnosing misdetections."""
    d = UPLOAD_DIR / session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")
    fields = extract_fields(str(cur))
    return {
        "headers": [
            {"paraIndex": f["paraIndex"], "text": f["text"][:60], "section": f["section"]["key"]}
            for f in fields if f.get("isHeader") and f.get("source") == "body"
        ],
        "all_fields": [
            {"id": f["id"], "paraIndex": f.get("paraIndex"), "section": f["section"]["key"],
             "isHeader": f.get("isHeader"), "text": f["text"][:60]}
            for f in fields if f.get("source") == "body"
        ][:60]
    }

@app.get("/debug/{session_id}")
async def debug(session_id: str, para: int=0, count: int=5):
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Not found")
    doc=Document(str(cur)); result=[]
    for i in range(para,min(para+count,len(doc.paragraphs))):
        p=doc.paragraphs[i]; ar=_get_runs(p); groups,_=_run_groups(p)
        result.append(dict(para=i,text=_para_text(p),
            runs=[dict(i=j,text=r["text"],bold=r["bold"]) for j,r in enumerate(ar)],
            groups=[dict(text=g["text"],indices=g["indices"]) for g in groups]))
    return dict(paragraphs=result)


@app.delete("/session/{session_id}")
async def del_session(session_id: str):
    d=UPLOAD_DIR/session_id
    if d.exists(): shutil.rmtree(d)
    return dict(success=True)


@app.get("/debug-para/{session_id}")
async def debug_para(session_id: str, para: int = 0):
    from lxml import etree as _et
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Not found")
    doc=Document(str(cur))
    p = doc.paragraphs[para]
    children = []
    for i, ch in enumerate(p._element):
        children.append({
            "index": i,
            "tag": ch.tag,
            "xml": _et.tostring(ch, pretty_print=True).decode()[:800]
        })
    return {"para_text": p.text, "num_children": len(children), "children": children}


@app.get("/debug-xml/{session_id}")
async def debug_xml(session_id: str, para: int = -1):
    d=UPLOAD_DIR/session_id; cur=d/"current.docx"
    if not cur.exists(): raise HTTPException(404,"Not found")
    doc=Document(str(cur))
    paras=doc.paragraphs
    idx = para if para >= 0 else len(paras)-1
    p = paras[idx]
    xml = etree.tostring(p._element, pretty_print=True).decode()
    import zipfile
    with zipfile.ZipFile(str(cur)) as z:
        styles = z.read('word/styles.xml').decode()
    import re
    style_colors = re.findall(r'<w:style[^>]*>.*?<w:color[^/]*/>', styles, re.DOTALL)[:5]
    return {"para_index": idx, "total": len(paras), "text": _para_text(p), "xml": xml, "style_colors": style_colors}


@app.get("/get-segments/{session_id}")
async def get_segments(session_id: str, para: int = -1, source: str = "body"):
    d = UPLOAD_DIR/session_id; cur = d/"current.docx"
    if not cur.exists(): raise HTTPException(404, "Session not found")
    try:
        doc = Document(str(cur))
        paras = _source_paras(doc, source)
        if not (0 <= para < len(paras)):
            return dict(success=False, segments=[])
        p = paras[para]
        runs = _get_runs(p)
        pPr = p._element.find(qn("w:pPr"))
        para_rpr = pPr.find(qn("w:rPr")) if pPr is not None else None

        def _rpr_val(rpr, tag):
            el = rpr.find(qn(tag)) if rpr is not None else None
            if el is None and para_rpr is not None:
                el = para_rpr.find(qn(tag))
            return el

        BULLET_CHARS_WIDE = {'•','◦','▪','▸','–','-','*','·','○','■','□','➢','✓','→'}
        segments = []
        found_text_seg = False
        for ri in runs:
            if ri.get("has_tab") and not ri["text"].strip():
                if found_text_seg:
                    segments.append(dict(
                        text="	", bold=False, italic=False,
                        underline=False, strike=False,
                        color="", fontFamily="", fontSize=""
                    ))
                continue
            if not ri["text"]: continue
            stripped = ri["text"].strip()
            if not found_text_seg and stripped and (
                (len(stripped) <= 2 and all(c in BULLET_CHARS_WIDE or ord(c)>127 for c in stripped))
            ):
                continue
            found_text_seg = True
            rpr = ri["elem"].find(qn("w:rPr"))
            bold = ri["bold"]
            i_el = _rpr_val(rpr, "w:i")
            italic = i_el is not None and i_el.get(qn("w:val"), "true") not in ("0","false")
            u_el = _rpr_val(rpr, "w:u")
            underline = u_el is not None and u_el.get(qn("w:val"),"single") not in ("none","0")
            s_el = _rpr_val(rpr, "w:strike")
            strike = s_el is not None
            color = ""
            c_el = _rpr_val(rpr, "w:color")
            if c_el is not None:
                val = c_el.get(qn("w:val"),"")
                if val and val.upper() not in ("000000","AUTO",""): color = "#"+val.upper()
            fontFamily = ""
            fe = _rpr_val(rpr, "w:rFonts")
            if fe is not None:
                fontFamily = fe.get(qn("w:ascii")) or fe.get(qn("w:hAnsi")) or ""
            fontSize = ""
            sz = _rpr_val(rpr, "w:sz")
            if sz is not None and sz.get(qn("w:val")):
                pt = int(sz.get(qn("w:val"))) / 2
                fontSize = f"{pt:.0f}pt"
            segments.append(dict(
                text=ri["text"], bold=bold, italic=italic,
                underline=underline, strike=strike, color=color,
                fontFamily=fontFamily, fontSize=fontSize
            ))
        merged = []
        for seg in segments:
            prev = merged[-1] if merged else None
            if (prev and prev["bold"]==seg["bold"] and prev["italic"]==seg["italic"]
                    and prev["underline"]==seg["underline"] and prev["strike"]==seg["strike"]
                    and prev["color"]==seg["color"] and prev["fontFamily"]==seg["fontFamily"]
                    and prev["fontSize"]==seg["fontSize"]):
                prev["text"] += seg["text"]
            else:
                merged.append(dict(seg))
        return dict(success=True, segments=merged)
    except Exception as e:
        return dict(success=False, segments=[], error=str(e))


@app.post("/add-page/{session_id}")
async def add_page(session_id: str):
    """
    Append a new full-width single-column page.

    Layout strategy (no extra paragraphs added to existing content):
      - If the body sectPr has multi-column layout:
          * Copy the original sectPr into the LAST EXISTING paragraph's pPr
            (as a continuous section break) so existing content keeps its layout.
          * Change the body sectPr to single-column — new content uses it.
          * No new paragraph is inserted before the page break — zero risk
            of pushing existing content (e.g. education) to the next page.
      - If already single-column: just add a plain page break + spacers.
      - Either way page_break_para_index points to the w:br paragraph.
    """
    d = UPLOAD_DIR/session_id; cur = d/"current.docx"
    if not cur.exists(): raise HTTPException(404, "Session not found")
    try:
        from docx.oxml import OxmlElement
        import copy

        doc  = Document(str(cur))
        body = doc.element.body

        body_sectPr = body.find(qn("w:sectPr"))

        # ── Detect multi-column layout ────────────────────────────────
        is_multicol = False
        if body_sectPr is not None:
            cols_el = body_sectPr.find(qn("w:cols"))
            if cols_el is not None:
                num = cols_el.get(qn("w:num"), "1")
                is_multicol = int(num) > 1

        if is_multicol and body_sectPr is not None:
            # Find the last existing body paragraph BEFORE we add anything
            existing_body_paras = [
                ch for ch in body if ch.tag == qn("w:p")
            ]
            if existing_body_paras:
                last_existing = existing_body_paras[-1]
                last_pPr = last_existing.find(qn("w:pPr"))
                if last_pPr is None:
                    last_pPr = last_existing.makeelement(qn("w:pPr"), {})
                    last_existing.insert(0, last_pPr)

                # If it already has a sectPr don't add another one
                if last_pPr.find(qn("w:sectPr")) is None:
                    # Copy original sectPr (with original cols) into last para
                    preserved_sectPr = copy.deepcopy(body_sectPr)
                    # Make it continuous so it doesn't force a page break
                    pt = preserved_sectPr.find(qn("w:type"))
                    if pt is None:
                        pt = preserved_sectPr.makeelement(qn("w:type"), {})
                        preserved_sectPr.insert(0, pt)
                    pt.set(qn("w:val"), "continuous")
                    last_pPr.append(preserved_sectPr)

            # Switch body sectPr to single column so new content is full-width
            cols_el = body_sectPr.find(qn("w:cols"))
            if cols_el is None:
                cols_el = body_sectPr.makeelement(qn("w:cols"), {})
                body_sectPr.append(cols_el)
            cols_el.set(qn("w:num"), "1")
            cols_el.set(qn("w:space"), "720")
            # Remove any individual col width definitions
            for col in list(cols_el):
                cols_el.remove(col)

        # ── page_break_para_index ─────────────────────────────────────
        page_break_para_index = len(doc.paragraphs)

        def _insert(elem):
            if body_sectPr is not None:
                body.insert(list(body).index(body_sectPr), elem)
            else:
                body.append(elem)

        # Page-break paragraph
        pb_para = OxmlElement("w:p")
        pb_run  = OxmlElement("w:r")
        pb_br   = OxmlElement("w:br")
        pb_br.set(qn("w:type"), "page")
        pb_run.append(pb_br)
        pb_para.append(pb_run)
        _insert(pb_para)

        # 20 blank spacer paragraphs for the new page
        for _ in range(20):
            _insert(OxmlElement("w:p"))

        tmp = d/"temp.docx"; doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        print(f"[AddPage] page at para {page_break_para_index}, multicol_fix={is_multicol}")
        return dict(
            success=True,
            fields=extract_fields(str(cur)),
            page_break_para_index=page_break_para_index,
        )
    except Exception as e:
        import traceback
        print(f"[AddPage] {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))

# ── AI Rewrite endpoint ──────────────────────────────────────────────
class AIRewriteRequest(BaseModel):
    text: str
    mode: str  # Polish, Elaborate, Formalise, Condense, Quantify, Action Verbs

@app.post("/ai-rewrite")
async def ai_rewrite(data: AIRewriteRequest):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(500, "OPENAI_API_KEY not configured")

    prompts = {
        "Polish":       f"Polish this resume bullet/text to be concise and professional. Return only the improved text, no explanation:\n\n{data.text}",
        "Elaborate":    f"Elaborate on this resume bullet to add more detail and impact. Return only the improved text, no explanation:\n\n{data.text}",
        "Formalise":    f"Rewrite this resume text in a formal, professional tone. Return only the improved text, no explanation:\n\n{data.text}",
        "Condense":     f"Condense this resume bullet to be shorter and punchier. Return only the improved text, no explanation:\n\n{data.text}",
        "Quantify":     f"Add quantifiable metrics/numbers to make this resume bullet more impactful. Infer reasonable numbers if needed. Return only the improved text, no explanation:\n\n{data.text}",
        "Action Verbs": f"Rewrite this resume bullet starting with a strong action verb. Return only the improved text, no explanation:\n\n{data.text}",
    }

    prompt = prompts.get(data.mode)
    if not prompt:
        raise HTTPException(400, f"Unknown mode: {data.mode}")

    response = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 500,
            "temperature": 0.7,
        },
        timeout=30,
    )

    if response.status_code != 200:
        raise HTTPException(500, f"OpenAI error: {response.text}")

    result = response.json()
    suggestion = result["choices"][0]["message"]["content"].strip()
    return {"suggestion": suggestion}



# ════════════════════════════════════════════════════════════════════
#  Move section to next page — inserts/removes w:pageBreakBefore
# ════════════════════════════════════════════════════════════════════

class MoveSectionRequest(BaseModel):
    session_id:  str
    para_index:  int   # first paragraph of the target section
    source:      str = "body"
    enable:      bool = True   # True = add page break, False = remove it

@app.post("/move-to-next-page")
async def move_to_next_page(req: MoveSectionRequest):
    """
    Toggle w:pageBreakBefore on the first paragraph of a section.
    enable=True  → section starts on a new page
    enable=False → remove the forced page break
    """
    d = UPLOAD_DIR / req.session_id
    cur = d / "current.docx"
    if not cur.exists():
        raise HTTPException(404, "Session not found")
    try:
        doc = Document(str(cur))
        paras = _source_paras(doc, req.source)
        if not (0 <= req.para_index < len(paras)):
            return dict(success=False, message=f"Para index {req.para_index} out of range")

        para = paras[req.para_index]
        p_elem = para._element

        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = p_elem.makeelement(qn("w:pPr"), {})
            p_elem.insert(0, pPr)

        pb = pPr.find(qn("w:pageBreakBefore"))
        if req.enable:
            if pb is None:
                pb = pPr.makeelement(qn("w:pageBreakBefore"), {})
                # Insert after w:jc / w:spacing / etc but before w:rPr
                rpr_pos = None
                for i, ch in enumerate(pPr):
                    if ch.tag == qn("w:rPr"):
                        rpr_pos = i; break
                if rpr_pos is not None:
                    pPr.insert(rpr_pos, pb)
                else:
                    pPr.append(pb)
            else:
                pb.attrib.pop(qn("w:val"), None)  # ensure it's not val="0"
        else:
            if pb is not None:
                pPr.remove(pb)

        tmp = d / "temp.docx"
        doc.save(str(tmp))
        shutil.move(str(tmp), str(cur))
        _patch_styles_gray(cur)

        action = "moved to next page" if req.enable else "page break removed"
        print(f"[MovePage] para {req.para_index} {action}")
        return dict(success=True, fields=extract_fields(str(cur)), page_break=req.enable)

    except Exception as e:
        import traceback
        print(f"[MovePage] {e}\n{traceback.format_exc()}")
        return dict(success=False, message=str(e))

@app.get("/health")
async def health():
    return dict(status="ok",version="7.6",
                converters=dict(ilovepdf=ILOVEPDF_OK,word_com=WORD_COM_OK,
                                pdf2docx=PDF2DOCX_OK,docx2pdf=DOCX2PDF_OK),
                ai=AI_SUPPORT)


if __name__=="__main__":
    import uvicorn
    uvicorn.run(app,host="0.0.0.0",port=8000)